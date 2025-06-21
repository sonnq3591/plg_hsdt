#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Walking Skeleton API - Phase 1
Simple FastAPI wrapper for your existing processors
Goal: Get ONE complete end-to-end flow working first
"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
import os
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import List
import uvicorn

# Import your existing processors
from processor import VietnameseProcurementProcessor
from processor_pham_vi import PhamViCungCapProcessor
from processor_can_cu import CanCuPhapLyProcessor
from processor_muc_dich import MucDichProcessor
from combined_processor import CombinedProcessor

app = FastAPI(title="Vietnamese Procurement Document API", version="1.0.0")

class WalkingSkeletonProcessor:
    """
    Walking Skeleton: Process ONE template completely
    Template: 02_MUC_DO_HIEU_BIET.docx
    Required placeholders: {{ten_goi_thau}}, {{pham_vi_cung_cap}}, {{can_cu_phap_ly}}, {{muc_dich_cong_viec}}, {{cac_buoc_thuc_hien}}
    """
    
    def __init__(self):
        self.work_dir = None
        self.pdf_inputs_dir = None
        
    def setup_workspace(self):
        """Create temporary workspace"""
        self.work_dir = Path(tempfile.mkdtemp())
        self.pdf_inputs_dir = self.work_dir / "pdf_inputs"
        self.pdf_inputs_dir.mkdir(exist_ok=True)
        print(f"🔧 Created workspace: {self.work_dir}")
        
    def save_uploaded_pdfs(self, pdf_files: List[UploadFile]):
        """Save uploaded PDFs to workspace"""
        expected_files = ['TBMT.pdf', 'BMMT.pdf', 'CHUONG_III.pdf', 'CHUONG_V.pdf', 'HSMT.pdf']
        saved_files = {}
        
        for pdf_file in pdf_files:
            if pdf_file.filename in expected_files:
                file_path = self.pdf_inputs_dir / pdf_file.filename
                with open(file_path, "wb") as f:
                    shutil.copyfileobj(pdf_file.file, f)
                saved_files[pdf_file.filename] = file_path
                print(f"✅ Saved: {pdf_file.filename}")
            else:
                print(f"⚠️ Unexpected file: {pdf_file.filename}")
        
        return saved_files
    
    def process_muc_do_hieu_biet_template(self):
        """
        FIXED: Call existing modular methods in sequence with proper file management
        Reuse all existing test_xxx methods, just control the template→output flow
        """
        print("\n🚀 WALKING SKELETON: Orchestrating Existing Modules")
        print("=" * 60)
        
        # Change working directory to our workspace
        original_cwd = os.getcwd()
        os.chdir(self.work_dir)
        
        try:
            # Copy template files to workspace
            template_source = Path(original_cwd) / "02_MUC_DO_HIEU_BIET_template.docx"
            template_dest = self.work_dir / "02_MUC_DO_HIEU_BIET_template.docx"
            shutil.copy2(template_source, template_dest)
            
            # Copy other required files (21_BUOC.docx, 23_BUOC.docx)
            for file_name in ["21_BUOC.docx", "23_BUOC.docx"]:
                source = Path(original_cwd) / file_name
                if source.exists():
                    shutil.copy2(source, self.work_dir / file_name)
            
            # PHASE 1: Call each module's existing test method (but manage output carefully)
            print("\n📋 PHASE 1: Call existing modular methods...")
            
            # Step 1: {{ten_goi_thau}} - let it create its output
            print("📋 Step 1: Calling processor.test_ten_goi_thau_extraction()...")
            processor1 = VietnameseProcurementProcessor()
            success1 = processor1.test_ten_goi_thau_extraction()
            if not success1:
                raise Exception("Failed {{ten_goi_thau}} module")
            
            # Backup the result (module creates "02_MUC_DO_HIEU_BIET_output.docx")
            shutil.copy2("02_MUC_DO_HIEU_BIET_output.docx", "step1_ten_goi_thau.docx")
            print("✅ Step 1 complete, backed up result")
            
            # Step 2: {{pham_vi_cung_cap}} - let it do its work
            print("📊 Step 2: Calling processor_pham_vi.test_pham_vi_cung_cap_simple()...")
            processor2 = PhamViCungCapProcessor()
            success2 = processor2.test_pham_vi_cung_cap_simple()
            if not success2:
                raise Exception("Failed {{pham_vi_cung_cap}} module")
            
            # Backup the result
            shutil.copy2("02_MUC_DO_HIEU_BIET_output.docx", "step2_pham_vi.docx")
            print("✅ Step 2 complete, backed up result")
            
            # Step 3: {{can_cu_phap_ly}} - let it do its work
            print("📜 Step 3: Calling processor_can_cu.test_can_cu_phap_ly_full_process()...")
            processor3 = CanCuPhapLyProcessor()
            success3 = processor3.test_can_cu_phap_ly_full_process()
            if not success3:
                raise Exception("Failed {{can_cu_phap_ly}} module")
            
            # Backup the result
            shutil.copy2("02_MUC_DO_HIEU_BIET_output.docx", "step3_can_cu.docx")
            print("✅ Step 3 complete, backed up result")
            
            # Step 4: {{muc_dich_cong_viec}} - let it do its work
            print("🎯 Step 4: Calling processor_muc_dich.test_muc_dich_cong_viec_full_process()...")
            processor4 = MucDichProcessor()
            success4 = processor4.test_muc_dich_cong_viec_full_process()
            if not success4:
                raise Exception("Failed {{muc_dich_cong_viec}} module")
            
            # Backup the result
            shutil.copy2("02_MUC_DO_HIEU_BIET_output.docx", "step4_muc_dich.docx")
            print("✅ Step 4 complete, backed up result")
            
            # PHASE 2: Now orchestrate the final combination
            print("\n🔄 PHASE 2: Orchestrate final combination...")
            
            # Start fresh with template
            shutil.copy2("02_MUC_DO_HIEU_BIET_template.docx", "02_MUC_DO_HIEU_BIET_output.docx")
            
            from docx import Document
            doc = Document("02_MUC_DO_HIEU_BIET_output.docx")
            
            # Apply each replacement using the modules' existing replace methods
            print("🔄 Applying {{ten_goi_thau}} from step1...")
            # Extract the content that was generated and apply it
            step1_doc = Document("step1_ten_goi_thau.docx")
            # Find what was replaced by comparing with template
            # For now, let's use the processor's replacement method
            processor1.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
            # Re-extract and apply (reuse the module's logic)
            tbmt_file = "pdf_inputs/TBMT.pdf"
            tbmt_content = processor1.extract_text_from_pdf(tbmt_file)
            ten_goi_thau_content = processor1.ask_openai_for_ten_goi_thau(tbmt_content)
            processor1.replace_placeholder_in_docx("{{ten_goi_thau}}", ten_goi_thau_content)
            print("✅ {{ten_goi_thau}} applied")
            
            print("🔄 Applying {{pham_vi_cung_cap}} from step2...")
            doc = Document("02_MUC_DO_HIEU_BIET_output.docx")
            processor2.replace_placeholder(doc, "{{pham_vi_cung_cap}}")
            doc.save("02_MUC_DO_HIEU_BIET_output.docx")
            print("✅ {{pham_vi_cung_cap}} applied")
            
            print("🔄 Applying {{can_cu_phap_ly}} from step3...")
            doc = Document("02_MUC_DO_HIEU_BIET_output.docx")
            processor3.replace_placeholder(doc, "{{can_cu_phap_ly}}")
            doc.save("02_MUC_DO_HIEU_BIET_output.docx")
            print("✅ {{can_cu_phap_ly}} applied")
            
            print("🔄 Applying {{muc_dich_cong_viec}} from step4...")
            doc = Document("02_MUC_DO_HIEU_BIET_output.docx")
            processor4.replace_placeholder(doc, "{{muc_dich_cong_viec}}")
            doc.save("02_MUC_DO_HIEU_BIET_output.docx")
            print("✅ {{muc_dich_cong_viec}} applied")
            
            # Step 5: {{cac_buoc_thuc_hien}} - let the module do its work LAST
            print("🔄 Step 5: Calling combined_processor.process_complete_workflow()...")
            processor5 = CombinedProcessor()
            success5 = processor5.process_complete_workflow()
            if not success5:
                raise Exception("Failed {{cac_buoc_thuc_hien}} module")
            
            print("✅ {{cac_buoc_thuc_hien}} applied (final step)")
            
            # Check if output file exists
            output_file = self.work_dir / "02_MUC_DO_HIEU_BIET_output.docx"
            if not output_file.exists():
                raise Exception("Output file not generated")
            
            print(f"\n🎉 MODULAR ORCHESTRATION SUCCESS!")
            print(f"📄 All existing modules called, all placeholders applied: {output_file}")
            return output_file
            
        except Exception as e:
            print(f"❌ Walking skeleton failed: {str(e)}")
            raise
        finally:
            os.chdir(original_cwd)
    
    def cleanup(self):
        """Clean up workspace"""
        if self.work_dir and self.work_dir.exists():
            shutil.rmtree(self.work_dir)
            print(f"🧹 Cleaned up workspace")

# API Endpoints

@app.post("/api/process-skeleton")
async def process_walking_skeleton(
    tbmt_pdf: UploadFile = File(..., description="TBMT.pdf file"),
    bmmt_pdf: UploadFile = File(..., description="BMMT.pdf file"),
    chuong_iii_pdf: UploadFile = File(..., description="CHUONG_III.pdf file"),
    chuong_v_pdf: UploadFile = File(..., description="CHUONG_V.pdf file"),
    hsmt_pdf: UploadFile = File(..., description="HSMT.pdf file")
):
    """
    Walking Skeleton Endpoint: Process ONE template completely
    Returns: 02_MUC_DO_HIEU_BIET_output.docx
    """
    processor = WalkingSkeletonProcessor()
    
    try:
        # Setup workspace
        processor.setup_workspace()
        
        # Save uploaded PDFs
        pdf_files = [tbmt_pdf, bmmt_pdf, chuong_iii_pdf, chuong_v_pdf, hsmt_pdf]
        saved_files = processor.save_uploaded_pdfs(pdf_files)
        
        # Validate all required PDFs are present
        required = {'TBMT.pdf', 'BMMT.pdf', 'CHUONG_III.pdf', 'CHUONG_V.pdf', 'HSMT.pdf'}
        missing = required - set(saved_files.keys())
        if missing:
            raise HTTPException(status_code=400, detail=f"Missing PDF files: {missing}")
        
        # Process template
        output_file = processor.process_muc_do_hieu_biet_template()
        
        # Return the generated DOCX file
        return FileResponse(
            path=output_file,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename="02_MUC_DO_HIEU_BIET_output.docx"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        processor.cleanup()

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "message": "Walking Skeleton API is running"}

@app.get("/api/templates")
async def list_templates():
    """List available templates (for future expansion)"""
    return {
        "available_templates": [
            "02_MUC_DO_HIEU_BIET"  # Only this one works in walking skeleton
        ],
        "future_templates": [
            "01_TINH_HIEU_QUA_CONG_VIEC",
            "03_CAC_BUOC_VA_CACH_THUC_THUC_HIEN",
            # ... other 12 templates
        ]
    }

if __name__ == "__main__":
    print("🚀 Starting Walking Skeleton API...")
    print("📋 Only 02_MUC_DO_HIEU_BIET template is implemented")
    print("🌐 API will be available at: http://localhost:8000")
    print("📖 API docs at: http://localhost:8000/docs")
    
    uvicorn.run(app, host="0.0.0.0", port=8000)