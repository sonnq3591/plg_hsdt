import os
import shutil
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt
import PyPDF2
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class VietnameseProcurementProcessor:
    def __init__(self):
        """Initialize the processor with OpenAI API key from .env"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        self.pdf_folder = Path("pdf_inputs")
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        
        # Create PDF folder if it doesn't exist
        self.pdf_folder.mkdir(exist_ok=True)
        
        # Expected PDF files
        self.pdf_files = {
            'TBMT': 'TBMT.pdf',
            'CHUONG_V': 'CHUONG_V.pdf', 
            'BMMT': 'BMMT.pdf',
            'CHUONG_III': 'CHUONG_III.pdf',
            'HSMT': 'HSMT.pdf'
        }
        
        print(f"📁 Created project structure:")
        print(f"   - PDF folder: {self.pdf_folder}")
        print(f"   - Template: {self.template_file}")
        print(f"   - Output: {self.output_file}")
        print(f"\n📋 Please place these PDFs in the '{self.pdf_folder}' folder:")
        for pdf in self.pdf_files.values():
            print(f"   - {pdf}")

    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"❌ Error reading {pdf_path}: {str(e)}")
            return None

    def ask_openai_for_ten_goi_thau(self, tbmt_content):
        """Ask OpenAI to extract 'ten_goi_thau' from TBMT.pdf content"""
        prompt = f"""
Bạn là chuyên gia phân tích tài liệu đấu thầu Việt Nam.

Từ nội dung tài liệu TBMT (Thông báo mời thầu) dưới đây, hãy tìm và trích xuất CHÍNH XÁC tên gói thầu.

HƯỚNG DẪN CỤ THỂ:
Tên gói thầu nằm trong bảng "Thông tin gói thầu" hoặc tương tự, tại dòng có:
- Cột trái: "Tên gói thầu" (có thể viết là "Tên dự án", "Tên gói", "Package name")
- Cột phải: [TÊN GÓI THẦU THỰC TẾ]

Các cách viết có thể gặp:
- "Tên gói thầu" | "Chỉnh lý tài liệu..."
- "Tên dự án" | "Chỉnh lý tài liệu..."  
- "Package name" | "Chỉnh lý tài liệu..."
- "Tên gói" | "Chỉnh lý tài liệu..."

YÊU CẦU:
1. CHỈ lấy nội dung từ cột bên PHẢI của dòng "Tên gói thầu"
2. KHÔNG lấy từ tiêu đề tài liệu hoặc nơi khác
3. KHÔNG bao gồm mã số gói thầu
4. Trích xuất CHÍNH XÁC, giữ nguyên dấu câu tiếng Việt
5. Trả về CHỈ tên gói thầu, không giải thích

ĐỊNH DẠNG MONG ĐỢI:
Tìm cấu trúc bảng dạng:
```
| Tên gói thầu | [TÊN CẦN TRÍCH XUẤT] |
```

NỘI DUNG TBMT:
{tbmt_content}

TÊN GÓI THẦU (chỉ nội dung cột phải):"""

        try:
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Bạn là chuyên gia trích xuất thông tin từ bảng trong tài liệu đấu thầu Việt Nam. Chỉ trả về nội dung được yêu cầu từ cột cụ thể trong bảng."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=300,
                temperature=0.0  # Even more deterministic
            )
            
            extracted_text = response.choices[0].message.content.strip()
            
            # Clean up any extra quotes or formatting
            if extracted_text.startswith('"') and extracted_text.endswith('"'):
                extracted_text = extracted_text[1:-1]
            
            print(f"🎯 OpenAI extracted from table: '{extracted_text}'")
            return extracted_text
            
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return "[KHÔNG TÌM THẤY]"

    def copy_template_to_output(self):
        """Copy template file to output file"""
        try:
            if not os.path.exists(self.template_file):
                print(f"❌ Template file not found: {self.template_file}")
                return False
            
            shutil.copy2(self.template_file, self.output_file)
            print(f"✅ Copied template to output file: {self.output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error copying template: {str(e)}")
            return False

    def replace_placeholder_in_docx(self, placeholder, content):
        """Replace placeholder in DOCX file while preserving exact formatting - simplified approach"""
        try:
            doc = Document(self.output_file)
            replaced = False
            
            print(f"🔍 Looking for placeholder: '{placeholder}'")
            print(f"🔄 Will replace with: '{content}'")
            
            # Simple approach: Replace in paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                full_text = paragraph.text
                if placeholder in full_text:
                    print(f"📍 Found placeholder in paragraph {para_idx}")
                    
                    # Build the replacement character by character to preserve formatting
                    placeholder_start = full_text.find(placeholder)
                    placeholder_end = placeholder_start + len(placeholder)
                    
                    # Find which runs contain the placeholder
                    char_count = 0
                    start_run_idx = -1
                    end_run_idx = -1
                    start_char_in_run = 0
                    end_char_in_run = 0
                    
                    for run_idx, run in enumerate(paragraph.runs):
                        run_len = len(run.text)
                        
                        # Check if placeholder starts in this run
                        if start_run_idx == -1 and char_count <= placeholder_start < char_count + run_len:
                            start_run_idx = run_idx
                            start_char_in_run = placeholder_start - char_count
                        
                        # Check if placeholder ends in this run
                        if char_count < placeholder_end <= char_count + run_len:
                            end_run_idx = run_idx
                            end_char_in_run = placeholder_end - char_count
                            break
                            
                        char_count += run_len
                    
                    if start_run_idx >= 0 and end_run_idx >= 0:
                        print(f"📍 Placeholder spans from run {start_run_idx} to run {end_run_idx}")
                        
                        # Case 1: Placeholder is within a single run
                        if start_run_idx == end_run_idx:
                            run = paragraph.runs[start_run_idx]
                            old_text = run.text
                            new_text = old_text[:start_char_in_run] + content + old_text[end_char_in_run:]
                            run.text = new_text
                            print(f"✅ Single run replacement: '{old_text}' → '{new_text}'")
                            
                        # Case 2: Placeholder spans multiple runs
                        else:
                            # Clear placeholder from all affected runs
                            for i in range(start_run_idx, end_run_idx + 1):
                                run = paragraph.runs[i]
                                if i == start_run_idx:
                                    # Keep text before placeholder
                                    run.text = run.text[:start_char_in_run]
                                elif i == end_run_idx:
                                    # Keep text after placeholder and add content
                                    run.text = content + run.text[end_char_in_run:]
                                else:
                                    # Clear middle runs
                                    run.text = ""
                            print(f"✅ Multi-run replacement completed")
                        
                        replaced = True
                        break  # Only replace first occurrence
            
            # Also check tables with same logic
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text = paragraph.text
                            if placeholder in full_text:
                                print(f"📍 Found placeholder in table cell")
                                
                                # Simple replacement for table cells
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, content)
                                        replaced = True
                                        break
            
            doc.save(self.output_file)
            
            if replaced:
                print(f"✅ Successfully replaced {placeholder}")
            else:
                print(f"❌ Failed to find {placeholder}")
            
            return replaced
            
        except Exception as e:
            print(f"❌ Error replacing placeholder: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def test_ten_goi_thau_extraction(self):
        """Test extraction of 'ten_goi_thau' from TBMT.pdf"""
        print("\n🧪 TESTING: {{ten_goi_thau}} extraction from TBMT.pdf")
        print("=" * 60)
        
        # Check if TBMT.pdf exists
        tbmt_file = self.pdf_folder / self.pdf_files['TBMT']
        if not tbmt_file.exists():
            print(f"❌ File not found: {tbmt_file}")
            print(f"📋 Please place TBMT.pdf in the '{self.pdf_folder}' folder")
            return False
        
        # Extract text from TBMT.pdf
        print("📖 Reading TBMT.pdf...")
        tbmt_content = self.extract_text_from_pdf(tbmt_file)
        if not tbmt_content:
            return False
        
        print(f"✅ Extracted {len(tbmt_content)} characters from TBMT.pdf")
        
        # Show preview of content
        preview = tbmt_content[:500] + "..." if len(tbmt_content) > 500 else tbmt_content
        print(f"📄 Content preview:\n{preview}\n")
        
        # Ask OpenAI to extract ten_goi_thau
        print("🤖 Asking OpenAI to extract 'ten_goi_thau'...")
        ten_goi_thau = self.ask_openai_for_ten_goi_thau(tbmt_content)
        
        print(f"📝 Extracted 'ten_goi_thau': {ten_goi_thau}")
        
        # Copy template and replace placeholder
        if self.copy_template_to_output():
            if self.replace_placeholder_in_docx("{{ten_goi_thau}}", ten_goi_thau):
                print(f"✅ SUCCESS: {{ten_goi_thau}} has been processed!")
                print(f"📄 Check output file: {self.output_file}")
                return True
        
        return False

def main():
    print("🇻🇳 Vietnamese Procurement Document Processor")
    print("=" * 50)
    
    try:
        # Initialize processor (API key loaded from .env)
        processor = VietnameseProcurementProcessor()
        print("✅ Environment loaded successfully")
        
        # Test ten_goi_thau extraction
        processor.test_ten_goi_thau_extraction()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")

if __name__ == "__main__":
    main()