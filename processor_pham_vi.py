import os
import shutil
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import PyPDF2
from dotenv import load_dotenv
import re
from copy import deepcopy
import csv
from io import StringIO

# Load environment variables
load_dotenv()

class PhamViCungCapProcessor:
    def __init__(self):
        """Initialize the processor for {{pham_vi_cung_cap}} - simple table extraction"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        
        self.pdf_folder = Path("pdf_inputs")
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        
        # Create processing folder structure
        self.process_folder = Path("processed/pham_vi_cung_cap")
        self.process_folder.mkdir(parents=True, exist_ok=True)
        
        print(f"🎯 PhamViCungCapProcessor initialized - simple table extraction")

    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"❌ Error reading {pdf_path}: {str(e)}")
            return None

    def extract_table_from_bmmt(self, bmmt_content):
        """Extract table data from BMMT content with specific column requirements"""
        prompt = f"""
From the BMMT document below, find the main service scope table and extract it with these EXACT columns only:

1. STT (sequential number)
2. Danh mục dịch vụ (service category/name)
3. Khối lượng (quantity/volume - KEEP EXACT FORMAT from PDF including commas/dots like "33,81" or "182,31")
4. Đơn vị tính (unit - like "Mét")  
5. Địa điểm thực hiện (implementation location - the full address)
6. Ngày hoàn thành (completion date - like "120 ngày")

CRITICAL FORMATTING RULES:
- For "Khối lượng" column: PRESERVE the exact number format from the original PDF
- If the PDF shows "33,81" keep it as "33,81" (with comma)
- If the PDF shows "182,31" keep it as "182,31" (with comma)
- Do NOT convert to decimals like "33.81" - keep the original Vietnamese number format
- SKIP the "Mô tả dịch vụ" column completely
- Keep all other Vietnamese text exactly as written

Use CSV format but put the Khối lượng numbers in quotes to preserve formatting:
STT,Danh mục dịch vụ,Khối lượng,Đơn vị tính,Địa điểm thực hiện,Ngày hoàn thành
1,"Service name 1","33,81",Mét,"Full address...","120 ngày"
2,"Service name 2","182,31",Mét,"Full address...","120 ngày"

BMMT CONTENT:
{bmmt_content}

CSV TABLE (preserve exact number formats):
"""

        try:
            response = openai.ChatCompletion.create(
                model='gpt-4o',
                messages=[
                    {"role": "system", "content": "You are an expert at extracting Vietnamese tables while preserving exact number formatting. Never change comma/dot formatting in numbers - keep them exactly as they appear in the original document."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.0
            )
            
            csv_data = response.choices[0].message.content.strip()
            
            # Clean up the CSV data
            lines = csv_data.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line and not line.startswith('```'):
                    cleaned_lines.append(line)
            
            final_csv = '\n'.join(cleaned_lines)
            print(f"🎯 OpenAI extracted table with {len(cleaned_lines)} rows (preserving number formats)")
            
            # Save for debugging
            (self.process_folder / 'extracted_csv.txt').write_text(final_csv, encoding='utf-8')
            
            # Verify column count and show sample
            if cleaned_lines:
                first_row = cleaned_lines[0]
                reader = csv.reader(StringIO(first_row))
                cols = next(reader)
                print(f"📊 Header columns: {cols}")
                if len(cols) != 6:
                    print(f"⚠️ Warning: Expected 6 columns, got {len(cols)}")
                
                # Show first data row to verify number formatting
                if len(cleaned_lines) > 1:
                    data_row = cleaned_lines[1]
                    reader = csv.reader(StringIO(data_row))
                    data = next(reader)
                    if len(data) >= 3:
                        print(f"📊 Sample Khối lượng value: '{data[2]}' (should preserve commas/dots)")
            
            return final_csv
            
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return None

    def create_docx_table(self, csv_data):
        """Convert CSV data to properly formatted DOCX table"""
        print(f"📊 Processing CSV data: {len(csv_data)} characters")
        
        lines = csv_data.strip().splitlines()
        print(f"📊 CSV has {len(lines)} lines")
        
        # Parse CSV with better handling
        parsed_rows = []
        for i, line in enumerate(lines):
            if line.strip():
                try:
                    reader = csv.reader(StringIO(line))
                    row = next(reader)
                    parsed_rows.append(row)
                    print(f"Row {i}: {len(row)} columns - {row}")
                except Exception as e:
                    print(f"⚠️ Error parsing line {i}: {e}")
                    continue

        if not parsed_rows:
            raise ValueError("No valid table data found")

        print(f"📊 Successfully parsed {len(parsed_rows)} rows")

        # Create DOCX with properly formatted table
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # Determine number of columns from first row
        num_cols = len(parsed_rows[0])
        print(f"📊 Table will have {num_cols} columns")

        # Calculate intelligent column widths
        max_lengths = [0] * num_cols
        for row in parsed_rows:
            for i in range(min(len(row), num_cols)):
                max_lengths[i] = max(max_lengths[i], len(str(row[i])))

        # Set column widths (total 7 inches)
        total_width = 7.0
        min_width = 0.6
        max_width = 2.0
        
        # Calculate proportional widths
        total_chars = sum(max_lengths) if sum(max_lengths) > 0 else num_cols
        column_widths = []
        
        for length in max_lengths:
            if total_chars > 0:
                width = (length / total_chars) * total_width
            else:
                width = total_width / num_cols
            width = max(min_width, min(width, max_width))
            column_widths.append(width)

        print(f"📊 Column widths: {[f'{w:.1f}' for w in column_widths]}")

        # Create table
        table = doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'
        
        # CENTER THE ENTIRE TABLE ON THE PAGE
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add rows with proper formatting
        for i, row_data in enumerate(parsed_rows):
            row = table.add_row().cells
            
            for j in range(num_cols):
                cell_text = str(row_data[j]) if j < len(row_data) else ""
                cell = row[j]
                
                # Set column width
                cell.width = Inches(column_widths[j])
                
                # Add text to cell
                paragraph = cell.paragraphs[0]
                paragraph.clear()  # Clear existing content
                run = paragraph.add_run(cell_text.strip())
                
                # FONT: 14pt Times New Roman for all content
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                
                # Header row formatting
                if i == 0:
                    run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Data row alignment based on content (keep same as before)
                    if j == 0 or len(cell_text.strip()) < 10:  # STT and short content
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:  # Long content
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Vertical alignment
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Save
        docx_path = self.process_folder / 'output.docx'
        doc.save(docx_path)
        print(f"✅ Created centered table with 14pt Times New Roman: {docx_path}")
        
        return docx_path

    def replace_placeholder(self, doc, placeholder_tag):
        """Replace placeholder using proven method with better debugging"""
        pattern = r"\{\{(.+?)\}\}"
        match = re.search(pattern, placeholder_tag)
        if not match:
            raise ValueError(f"Invalid placeholder: {placeholder_tag}")
        
        tag_name = match.group(1)
        source_path = Path(f"processed/{tag_name}/output.docx")

        print(f"🔍 Looking for source file: {source_path}")
        if not source_path.exists():
            raise FileNotFoundError(f"Missing source doc: {source_path}")

        source_doc = Document(source_path)
        
        # Debug: Check what's in the source document
        print(f"📊 Source document has {len(source_doc.paragraphs)} paragraphs")
        print(f"📊 Source document has {len(source_doc.tables)} tables")
        
        # Get all content (paragraphs and tables)
        all_elements = []
        
        # Add paragraphs
        for para in source_doc.paragraphs:
            if para.text.strip():
                all_elements.append(('paragraph', para))
        
        # Add tables
        for table in source_doc.tables:
            all_elements.append(('table', table))
        
        print(f"📊 Total elements to copy: {len(all_elements)}")

        # Find and replace placeholder
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                print(f"📍 Found placeholder in paragraph {i}")
                
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)
                
                print(f"🔄 Copying {len(all_elements)} elements...")

                # Insert all elements from source
                for element_type, element in reversed(all_elements):
                    if element_type == 'paragraph':
                        new_p = deepcopy(element._element)
                        parent.insert(index, new_p)
                    elif element_type == 'table':
                        new_t = deepcopy(element._element)
                        parent.insert(index, new_t)
                
                print(f"✅ Successfully copied all elements")
                break

    def copy_template_to_output(self):
        """Copy template file to output file"""
        try:
            if not os.path.exists(self.template_file):
                print(f"❌ Template file not found: {self.template_file}")
                return False
            
            shutil.copy2(self.template_file, self.output_file)
            print(f"✅ Copied template to output file: {self.output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error copying template: {str(e)}")
            return False

    def test_pham_vi_cung_cap_simple(self):
        """Simple table extraction and processing"""
        print("\n🧪 TESTING: {{pham_vi_cung_cap}} - Simple Table Extraction")
        print("=" * 60)
        
        # Check if BMMT.pdf exists
        bmmt_file = self.pdf_folder / 'BMMT.pdf'
        if not bmmt_file.exists():
            print(f"❌ File not found: {bmmt_file}")
            return False
        
        # Step 1: Extract text from PDF
        print("📖 Step 1: Reading BMMT.pdf...")
        bmmt_content = self.extract_text_from_pdf(bmmt_file)
        if not bmmt_content:
            return False
        
        print(f"✅ Extracted {len(bmmt_content)} characters")
        
        # Step 2: Extract table as CSV
        print("📊 Step 2: Extracting table from BMMT...")
        csv_data = self.extract_table_from_bmmt(bmmt_content)
        
        if not csv_data:
            print("❌ Failed to extract table")
            return False
        
        # Save CSV data
        (self.process_folder / 'input.txt').write_text(csv_data, encoding='utf-8')
        print(f"✅ Saved CSV data")
        
        # Step 3: Create DOCX table
        print("🔄 Step 3: Creating formatted table...")
        try:
            docx_path = self.create_docx_table(csv_data)
        except Exception as e:
            print(f"❌ Failed to create table: {str(e)}")
            return False
        
        # Step 4: Replace placeholder
        print("🔄 Step 4: Replacing placeholder in template...")
        if self.copy_template_to_output():
            doc = Document(self.output_file)
            self.replace_placeholder(doc, "{{pham_vi_cung_cap}}")
            doc.save(self.output_file)
            print(f"✅ SUCCESS: {{pham_vi_cung_cap}} table processed!")
            print(f"📄 Check output file: {self.output_file}")
            return True
        
        return False

def main():
    print("🇻🇳 Pham Vi Cung Cap Processor - Simple Table Extraction")
    print("=" * 60)
    
    try:
        processor = PhamViCungCapProcessor()
        print("✅ Environment loaded successfully")
        
        processor.test_pham_vi_cung_cap_simple()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()