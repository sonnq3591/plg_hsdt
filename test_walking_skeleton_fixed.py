#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Test Walking Skeleton Locally - FIXED VERSION
This tests your complete workflow using existing modular methods
"""

import os
import shutil
from pathlib import Path

class LocalWalkingSkeletonTest:
    """Test the complete walking skeleton workflow locally"""
    
    def __init__(self):
        self.required_files = [
            "02_MUC_DO_HIEU_BIET_template.docx",
            "21_BUOC.docx", 
            "23_BUOC.docx",
            "pdf_inputs/TBMT.pdf",
            "pdf_inputs/BMMT.pdf", 
            "pdf_inputs/CHUONG_III.pdf",
            "pdf_inputs/CHUONG_V.pdf",
            "pdf_inputs/HSMT.pdf"
        ]
        
    def check_required_files(self):
        """Check if all required files exist"""
        print("🔍 CHECKING REQUIRED FILES...")
        print("=" * 50)
        
        missing_files = []
        for file_path in self.required_files:
            if not Path(file_path).exists():
                missing_files.append(file_path)
                print(f"❌ Missing: {file_path}")
            else:
                print(f"✅ Found: {file_path}")
        
        if missing_files:
            print(f"\n❌ MISSING {len(missing_files)} required files!")
            print("📋 Please ensure these files exist before running the test:")
            for f in missing_files:
                print(f"   - {f}")
            return False
        
        print(f"\n✅ All {len(self.required_files)} required files found!")
        return True
    
    def clean_previous_outputs(self):
        """Clean up any previous outputs"""
        print("\n🧹 CLEANING PREVIOUS OUTPUTS...")
        
        # Clean main output
        output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        if Path(output_file).exists():
            os.remove(output_file)
            print(f"🗑️ Removed: {output_file}")
        
        # Clean processed folders
        processed_dir = Path("processed")
        if processed_dir.exists():
            shutil.rmtree(processed_dir)
            print(f"🗑️ Removed: {processed_dir}")
        
        # Clean backup files
        for backup_file in ["step1_ten_goi_thau.docx", "step2_pham_vi.docx", "step3_can_cu.docx", "step4_muc_dich.docx"]:
            if Path(backup_file).exists():
                os.remove(backup_file)
                print(f"🗑️ Removed: {backup_file}")
        
        print("✅ Cleanup complete")
    
    def test_complete_workflow(self):
        """Test using existing modular methods - sequential accumulation"""
        print("\n🚀 TESTING COMPLETE WALKING SKELETON WORKFLOW")
        print("=" * 60)
        print("🎯 Target: Generate 02_MUC_DO_HIEU_BIET_output.docx")
        print("📋 Each module works independently, accumulate results sequentially")
        print("=" * 60)
        
        try:
            # Step 1: {{ten_goi_thau}} - let module do its complete work
            print("📋 Step 1/5: Calling processor.test_ten_goi_thau_extraction()...")
            from processor import VietnameseProcurementProcessor
            processor1 = VietnameseProcurementProcessor()
            success1 = processor1.test_ten_goi_thau_extraction()
            if not success1:
                raise Exception("❌ STEP 1 FAILED: {{ten_goi_thau}}")
            print("✅ Step 1 SUCCESS: {{ten_goi_thau}} processed")
            
            # Step 2: {{pham_vi_cung_cap}} - let module work on the existing output
            print("📊 Step 2/5: Calling processor_pham_vi.test_pham_vi_cung_cap_simple()...")
            from processor_pham_vi import PhamViCungCapProcessor
            processor2 = PhamViCungCapProcessor()
            # Make sure it works on the current output (which has {{ten_goi_thau}} already replaced)
            processor2.template_file = "02_MUC_DO_HIEU_BIET_output.docx"  # Use current output as template
            processor2.output_file = "02_MUC_DO_HIEU_BIET_output.docx"    # Save to same file
            
            # Override the copy method to handle same file case
            original_copy_method = processor2.copy_template_to_output
            def smart_copy():
                if processor2.template_file == processor2.output_file:
                    print(f"✅ Template and output are same file - no copy needed")
                    return True
                return original_copy_method()
            processor2.copy_template_to_output = smart_copy
            
            success2 = processor2.test_pham_vi_cung_cap_simple()
            if not success2:
                raise Exception("❌ STEP 2 FAILED: {{pham_vi_cung_cap}}")
            print("✅ Step 2 SUCCESS: {{pham_vi_cung_cap}} processed")
            
            # Step 3: {{can_cu_phap_ly}} - let module work on the existing output
            print("📜 Step 3/5: Calling processor_can_cu.test_can_cu_phap_ly_full_process()...")
            from processor_can_cu import CanCuPhapLyProcessor
            processor3 = CanCuPhapLyProcessor()
            # Make sure it works on the current output (which has {{ten_goi_thau}} + {{pham_vi_cung_cap}} already replaced)
            processor3.template_file = "02_MUC_DO_HIEU_BIET_output.docx"  # Use current output as template
            processor3.output_file = "02_MUC_DO_HIEU_BIET_output.docx"    # Save to same file
            
            # Override the copy method to handle same file case
            original_copy_method3 = processor3.copy_template_to_output
            def smart_copy3():
                if processor3.template_file == processor3.output_file:
                    print(f"✅ Template and output are same file - no copy needed")
                    return True
                return original_copy_method3()
            processor3.copy_template_to_output = smart_copy3
            
            success3 = processor3.test_can_cu_phap_ly_full_process()
            if not success3:
                raise Exception("❌ STEP 3 FAILED: {{can_cu_phap_ly}}")
            print("✅ Step 3 SUCCESS: {{can_cu_phap_ly}} processed")
            
            # Step 4: {{muc_dich_cong_viec}} - let module work on the existing output
            print("🎯 Step 4/5: Calling processor_muc_dich.test_muc_dich_cong_viec_full_process()...")
            from processor_muc_dich import MucDichProcessor
            processor4 = MucDichProcessor()
            # Make sure it works on the current output (which has 3 placeholders already replaced)
            processor4.template_file = "02_MUC_DO_HIEU_BIET_output.docx"  # Use current output as template
            processor4.output_file = "02_MUC_DO_HIEU_BIET_output.docx"    # Save to same file
            
            # Override the copy method to handle same file case
            original_copy_method4 = processor4.copy_template_to_output
            def smart_copy4():
                if processor4.template_file == processor4.output_file:
                    print(f"✅ Template and output are same file - no copy needed")
                    return True
                return original_copy_method4()
            processor4.copy_template_to_output = smart_copy4
            
            success4 = processor4.test_muc_dich_cong_viec_full_process()
            if not success4:
                raise Exception("❌ STEP 4 FAILED: {{muc_dich_cong_viec}}")
            print("✅ Step 4 SUCCESS: {{muc_dich_cong_viec}} processed")
            
            # Step 5: {{cac_buoc_thuc_hien}} - let module work on the existing output
            print("🔄 Step 5/5: Calling combined_processor.process_complete_workflow()...")
            from combined_processor import CombinedProcessor
            processor5 = CombinedProcessor()
            # Make sure it works on the current output (which has 4 placeholders already replaced)
            processor5.template_file = "02_MUC_DO_HIEU_BIET_output.docx"  # Use current output as template
            processor5.output_file = "02_MUC_DO_HIEU_BIET_output.docx"    # Save to same file
            
            # Override the replace_placeholder_only method to handle same file case
            original_replace_method = processor5.replace_placeholder_only
            def smart_replace(template_path, content_path, output_path, placeholder="{{cac_buoc_thuc_hien}}"):
                # If template and output are same, don't copy - work directly on the file
                if template_path == output_path:
                    print(f"🔄 Working directly on file (no copy needed): {output_path}")
                    from docx import Document
                    from copy import deepcopy
                    
                    # Load the current document and content document
                    doc = Document(output_path)
                    source_doc = Document(content_path)
                    
                    # Get all content from source
                    all_elements = []
                    
                    for para in source_doc.paragraphs:
                        if para.text.strip():
                            all_elements.append(('paragraph', para))
                    
                    for table in source_doc.tables:
                        all_elements.append(('table', table))
                    
                    print(f"📊 Total elements to copy: {len(all_elements)}")
                    
                    # Find and replace placeholder
                    placeholder_found = False
                    
                    for i, paragraph in enumerate(doc.paragraphs):
                        if placeholder in paragraph.text:
                            print(f"📍 Found placeholder in paragraph {i}")
                            placeholder_found = True
                            
                            # Use the existing replacement logic
                            p_element = paragraph._element
                            parent = p_element.getparent()
                            index = parent.index(p_element)
                            parent.remove(p_element)
                            
                            # Insert all elements
                            for element_type, element in reversed(all_elements):
                                if element_type == 'paragraph':
                                    new_p = deepcopy(element._element)
                                    parent.insert(index, new_p)
                                elif element_type == 'table':
                                    new_t = deepcopy(element._element)
                                    parent.insert(index, new_t)
                            
                            break
                    
                    if not placeholder_found:
                        print(f"❌ Placeholder '{placeholder}' not found!")
                        return False
                    
                    # Save document
                    doc.save(output_path)
                    print(f"✅ Document saved: {output_path}")
                    return True
                else:
                    # Different files - use original method
                    return original_replace_method(template_path, content_path, output_path, placeholder)
            
            processor5.replace_placeholder_only = smart_replace
            
            success5 = processor5.process_complete_workflow()
            if not success5:
                raise Exception("❌ STEP 5 FAILED: {{cac_buoc_thuc_hien}}")
            print("✅ Step 5 SUCCESS: {{cac_buoc_thuc_hien}} processed")
            
            # Final validation
            output_file = Path("02_MUC_DO_HIEU_BIET_output.docx")
            if not output_file.exists():
                raise Exception("❌ FINAL OUTPUT FILE NOT FOUND!")
            
            file_size = output_file.stat().st_size
            print(f"\n🎉 SEQUENTIAL ACCUMULATION SUCCESS! 🎉")
            print("=" * 60)
            print(f"📄 Generated: {output_file}")
            print(f"📏 File size: {file_size:,} bytes")
            print(f"✅ All 5 placeholders replaced sequentially:")
            print(f"   1. {{{{ten_goi_thau}}}} ✅ (Step 1)")
            print(f"   2. {{{{pham_vi_cung_cap}}}} ✅ (Step 2 on Step 1 result)")
            print(f"   3. {{{{can_cu_phap_ly}}}} ✅ (Step 3 on Step 2 result)") 
            print(f"   4. {{{{muc_dich_cong_viec}}}} ✅ (Step 4 on Step 3 result)")
            print(f"   5. {{{{cac_buoc_thuc_hien}}}} ✅ (Step 5 on Step 4 result)")
            print("📋 Each module worked on the accumulated result!")
            print("=" * 60)
            print("🚀 READY FOR API INTEGRATION!")
            
            return True
            
        except Exception as e:
            print(f"\n💥 WALKING SKELETON FAILED!")
            print(f"❌ Error: {str(e)}")
            print("=" * 60)
            return False
    
    def show_next_steps(self):
        """Show next steps for walking skeleton"""
        print("\n📋 NEXT STEPS FOR WALKING SKELETON:")
        print("=" * 50)
        print("1. ✅ Local workflow test (this script)")
        print("2. 🔄 Wrap in FastAPI endpoint") 
        print("3. 🤖 Create basic Teams bot")
        print("4. 📤 Add SharePoint upload")
        print("5. 🔢 Add remaining 14 templates")
        print("\n💡 Current milestone: Get ONE template working end-to-end")

def main():
    """Main test execution"""
    print("🇻🇳 WALKING SKELETON LOCAL TEST")
    print("🎯 Testing: 02_MUC_DO_HIEU_BIET.docx generation")
    print("=" * 60)
    
    tester = LocalWalkingSkeletonTest()
    
    # Step 1: Check required files
    if not tester.check_required_files():
        print("\n🛑 Cannot proceed without required files!")
        return False
    
    # Step 2: Clean previous outputs
    tester.clean_previous_outputs()
    
    # Step 3: Test complete workflow
    success = tester.test_complete_workflow()
    
    # Step 4: Show next steps
    tester.show_next_steps()
    
    return success

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)