#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Combined Processor: Step Detection + Document Replacement
1. Detect steps (21 or 23) from CHUONG_V.pdf using OpenAI
2. Replace {{cac_buoc_thuc_hien}} with appropriate content
"""

import os
import shutil
import fitz  # PyMuPDF
import openai
from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class CombinedProcessor:
    def __init__(self):
        """Initialize the combined processor"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        
        # File paths
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        self.chuong_v_pdf = "CHUONG_V.pdf"
        
        print("✅ CombinedProcessor initialized")

    # ============ STEP DETECTION METHODS ============
    
    def extract_pdf_text_pymupdf(self, pdf_path):
        """Extract text from PDF using PyMuPDF"""
        try:
            print(f"📖 Extracting text from: {pdf_path}")
            
            doc = fitz.open(pdf_path)
            text = ""
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += f"\n--- PAGE {page_num + 1} ---\n"
                text += page_text
            
            doc.close()
            
            print(f"✅ Extracted {len(text)} characters from {page_count} pages")
            return text
            
        except Exception as e:
            print(f"❌ Error extracting PDF text: {str(e)}")
            return None

    def detect_steps_with_openai(self, pdf_text):
        """Ask OpenAI to analyze PDF content and determine step count"""
        
        prompt = f"""
Analyze this Vietnamese document content from CHUONG_V.pdf and find the section about implementation steps.

Look for sections with these characteristics:
1. Title containing "quy trình chỉnh lý" (implementation process) 
2. Content mentioning "Tuân thủ theo các bước thực hiện chỉnh lý" (follow implementation steps)
3. Reference to "trình tự 21 bước công việc" (21-step process) OR "trình tự 23 bước công việc" (23-step process)
4. A table with "Số TT" (sequential number) and "Nội dung công việc" (work content) columns
5. Section 3.2 or similar numbering about "Yêu cầu về quy trình chỉnh lý" or "Công việc thực hiện của mỗi bước"

Your task:
- Find the section that describes the step-by-step implementation process
- Count if it mentions 21 steps or 23 steps in the process
- Look for phrases like "21 bước công việc", "23 bước công việc", or count the actual steps in any process table

Return ONLY the number: "21" or "23"
If you cannot determine clearly, return "UNKNOWN"

DOCUMENT CONTENT:
{pdf_text}

STEP COUNT:"""

        try:
            print("🤖 Asking OpenAI to analyze step count...")
            
            response = openai.ChatCompletion.create(
                model='gpt-4o',
                messages=[
                    {
                        "role": "system", 
                        "content": "You are an expert at analyzing Vietnamese documents about administrative processes. You specialize in finding implementation step counts in document sections. Be precise and only return the step count number."
                    },
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ],
                max_tokens=50,
                temperature=0.0
            )
            
            step_count = response.choices[0].message.content.strip()
            print(f"🤖 OpenAI response: '{step_count}'")
            
            if step_count in ["21", "23"]:
                return int(step_count)
            elif step_count == "UNKNOWN":
                print("⚠️ OpenAI couldn't determine step count")
                return None
            else:
                print(f"⚠️ Unexpected response: {step_count}")
                return None
                
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return None

    # ============ DOCUMENT REPLACEMENT METHODS ============
    
    def replace_placeholder_only(self, template_path, content_path, output_path, placeholder="{{cac_buoc_thuc_hien}}"):
        """Replace ONLY the placeholder with content from another DOCX file"""
        try:
            print(f"🔄 Replacing placeholder with content from: {content_path}")
            
            # Step 1: Copy template to output
            shutil.copy2(template_path, output_path)
            
            # Step 2: Load documents
            doc = Document(output_path)
            source_doc = Document(content_path)
            
            # Step 3: Get all content from source
            all_elements = []
            
            for para in source_doc.paragraphs:
                if para.text.strip():
                    all_elements.append(('paragraph', para))
            
            for table in source_doc.tables:
                all_elements.append(('table', table))
            
            print(f"📊 Total elements to copy: {len(all_elements)}")
            
            # Step 4: Find and replace placeholder
            placeholder_found = False
            
            for i, paragraph in enumerate(doc.paragraphs):
                if placeholder in paragraph.text:
                    print(f"📍 Found placeholder in paragraph {i}")
                    placeholder_found = True
                    
                    # Use proven replacement method
                    p_element = paragraph._element
                    parent = p_element.getparent()
                    index = parent.index(p_element)
                    parent.remove(p_element)
                    
                    # Insert all elements with formatting enhancements
                    for element_type, element in reversed(all_elements):
                        if element_type == 'paragraph':
                            new_p = deepcopy(element._element)
                            
                            # ENHANCE PARAGRAPH SPACING
                            pPr = new_p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                            if pPr is None:
                                pPr = parse_xml('<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                                new_p.insert(0, pPr)
                            
                            spacing_xml = '''<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                            w:before="120" w:after="120" w:line="360" w:lineRule="auto"/>'''
                            spacing = parse_xml(spacing_xml)
                            
                            existing_spacing = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                            if existing_spacing is not None:
                                pPr.remove(existing_spacing)
                            pPr.append(spacing)
                            
                            parent.insert(index, new_p)
                            
                        elif element_type == 'table':
                            new_t = deepcopy(element._element)
                            
                            # ENHANCE TABLE ROW HEIGHT
                            rows = new_t.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
                            for row in rows:
                                trPr = row.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr')
                                if trPr is None:
                                    trPr = parse_xml('<w:trPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                                    row.insert(0, trPr)
                                
                                height_xml = '''<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                               w:val="600" w:hRule="atLeast"/>'''
                                height = parse_xml(height_xml)
                                
                                existing_height = trPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trHeight')
                                if existing_height is not None:
                                    trPr.remove(existing_height)
                                trPr.append(height)
                            
                            parent.insert(index, new_t)
                    
                    break
            
            if not placeholder_found:
                print(f"❌ Placeholder '{placeholder}' not found!")
                return False
            
            # Step 5: Save document
            doc.save(output_path)
            print(f"✅ Document saved: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error in replacement: {str(e)}")
            return False

    # ============ COMBINED PROCESS ============
    
    def process_complete_workflow(self):
        """Complete workflow: detect steps + replace content"""
        print("\n🚀 STARTING COMPLETE WORKFLOW")
        print("=" * 60)
        
        # Check required files
        required_files = [self.template_file, self.chuong_v_pdf, "21_BUOC.docx", "23_BUOC.docx"]
        missing_files = [f for f in required_files if not os.path.exists(f)]
        
        if missing_files:
            print("❌ Missing required files:")
            for f in missing_files:
                print(f"  - {f}")
            return False
        
        # Step 1: Extract PDF text
        print("\n📖 STEP 1: Extracting text from CHUONG_V.pdf...")
        pdf_text = self.extract_pdf_text_pymupdf(self.chuong_v_pdf)
        
        if not pdf_text:
            print("❌ Failed to extract PDF text")
            return False
        
        # Step 2: Detect step count
        print("\n🤖 STEP 2: Detecting step count with OpenAI...")
        step_count = self.detect_steps_with_openai(pdf_text)
        
        if not step_count:
            print("❌ Failed to detect step count")
            return False
        
        print(f"✅ Detected: {step_count} steps")
        
        # Step 3: Select appropriate content file
        content_file = f"{step_count}_BUOC.docx"
        print(f"\n📄 STEP 3: Selected content file: {content_file}")
        
        # Step 4: Replace placeholder
        print(f"\n🔄 STEP 4: Replacing placeholder with {step_count}-step content...")
        success = self.replace_placeholder_only(
            self.template_file, 
            content_file, 
            self.output_file
        )
        
        if success:
            print(f"\n🎉 SUCCESS! Complete workflow finished!")
            print(f"📄 Output file: {self.output_file}")
            print(f"🔢 Used {step_count}-step process")
            print("=" * 60)
            return True
        else:
            print(f"\n❌ FAILED during replacement step")
            return False

def main():
    """Main execution function"""
    print("🇻🇳 Combined Step Detection + Document Replacement")
    print("=" * 60)
    
    try:
        processor = CombinedProcessor()
        processor.process_complete_workflow()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()