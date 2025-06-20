#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Simple DOCX Placeholder Replacement
Following the proven method from processor_cac_buoc_premade.py
ONLY replaces the placeholder, doesn't touch anything else!
"""

import os
import shutil
from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy

def replace_placeholder_only(template_path, content_path, output_path, placeholder="{{cac_buoc_thuc_hien}}"):
    """
    Replace ONLY the placeholder with content from another DOCX file
    Following the exact proven method - don't touch anything else!
    """
    print("=== Simple Placeholder Replacement ===")
    print(f"Template: {template_path}")
    print(f"Content: {content_path}")
    print(f"Output: {output_path}")
    print(f"Placeholder: {placeholder}")
    print("-" * 50)
    
    try:
        # Step 1: Copy template to output (don't modify original)
        print("📋 Step 1: Copying template to output...")
        shutil.copy2(template_path, output_path)
        print(f"✅ Template copied to: {output_path}")
        
        # Step 2: Load the documents
        print("📖 Step 2: Loading documents...")
        doc = Document(output_path)  # The output document to modify
        source_doc = Document(content_path)  # The content to insert
        
        print(f"✅ Output doc loaded: {len(doc.paragraphs)} paragraphs")
        print(f"✅ Source doc loaded: {len(source_doc.paragraphs)} paragraphs, {len(source_doc.tables)} tables")
        
        # Step 3: Get all content from source (paragraphs and tables)
        print("📊 Step 3: Preparing content to copy...")
        all_elements = []
        
        # Add paragraphs that have content
        for para in source_doc.paragraphs:
            if para.text.strip():  # Only non-empty paragraphs
                all_elements.append(('paragraph', para))
        
        # Add all tables
        for table in source_doc.tables:
            all_elements.append(('table', table))
        
        print(f"📊 Total elements to copy: {len(all_elements)}")
        
        # Step 4: Find and replace placeholder (EXACT METHOD from working code)
        print("🔍 Step 4: Finding and replacing placeholder...")
        placeholder_found = False
        
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                print(f"📍 Found placeholder in paragraph {i}: '{paragraph.text[:50]}...'")
                placeholder_found = True
                
                # EXACT method from working code
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)
                
                print(f"🔄 Inserting {len(all_elements)} elements...")
                
                # Insert all elements from source (in reverse order to maintain position)
                for element_type, element in reversed(all_elements):
                    if element_type == 'paragraph':
                        new_p = deepcopy(element._element)
                        
                        # ENHANCE PARAGRAPH SPACING
                        # Add spacing before and after paragraph
                        pPr = new_p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                        if pPr is None:
                            pPr = parse_xml('<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                            new_p.insert(0, pPr)
                        
                        # Set spacing before and after (in twentieths of a point)
                        spacing_xml = '''<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                        w:before="120" w:after="120" w:line="360" w:lineRule="auto"/>'''
                        spacing = parse_xml(spacing_xml)
                        
                        # Remove existing spacing if any
                        existing_spacing = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                        if existing_spacing is not None:
                            pPr.remove(existing_spacing)
                        pPr.append(spacing)
                        
                        parent.insert(index, new_p)
                        
                    elif element_type == 'table':
                        new_t = deepcopy(element._element)
                        
                        # ENHANCE TABLE ROW HEIGHT
                        # Find all rows in the table
                        rows = new_t.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
                        for row in rows:
                            # Find or create trPr (table row properties)
                            trPr = row.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr')
                            if trPr is None:
                                trPr = parse_xml('<w:trPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                                row.insert(0, trPr)
                            
                            # Set minimum row height (in twentieths of a point)
                            # 600 = 30pt row height
                            height_xml = '''<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                           w:val="600" w:hRule="atLeast"/>'''
                            height = parse_xml(height_xml)
                            
                            # Remove existing height if any
                            existing_height = trPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trHeight')
                            if existing_height is not None:
                                trPr.remove(existing_height)
                            trPr.append(height)
                        
                        parent.insert(index, new_t)
                
                print(f"✅ Successfully replaced placeholder with all content")
                break
        
        if not placeholder_found:
            print(f"❌ Placeholder '{placeholder}' not found in document!")
            return False
        
        # Step 5: Save the modified document
        print("💾 Step 5: Saving final document...")
        doc.save(output_path)
        print(f"✅ Document saved: {output_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main execution - simple and focused"""
    # File paths
    template_file = "02_MUC_DO_HIEU_BIET_template.docx"
    content_file = "21_BUOC.docx"
    output_file = "02_MUC_DO_HIEU_BIET_output.docx"
    
    # Check if required files exist
    missing_files = []
    for file_path in [template_file, content_file]:
        if not os.path.exists(file_path):
            missing_files.append(file_path)
    
    if missing_files:
        print("❌ Missing required files:")
        for file_path in missing_files:
            print(f"  - {file_path}")
        return False
    
    # Do the replacement
    success = replace_placeholder_only(template_file, content_file, output_file)
    
    if success:
        print("\n" + "=" * 50)
        print("🎉 SUCCESS!")
        print(f"📄 Placeholder replaced successfully!")
        print(f"📁 Output file: {output_file}")
        print("🔹 Only the placeholder was replaced - nothing else touched!")
        print("=" * 50)
    else:
        print("\n" + "=" * 50)
        print("❌ FAILED!")
        print("Check error messages above.")
        print("=" * 50)
    
    return success

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)