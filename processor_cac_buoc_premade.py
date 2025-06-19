import os
import shutil
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt
import PyPDF2
from dotenv import load_dotenv
import re
from copy import deepcopy

# Load environment variables
load_dotenv()

class CacBuocThucHienProcessor:
    def __init__(self):
        """Initialize the processor for {{cac_buoc_thuc_hien}} with step counting logic"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        
        self.pdf_folder = Path("pdf_inputs")
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        
        # Create processing folder structure
        self.process_folder = Path("processed/cac_buoc_thuc_hien")
        self.process_folder.mkdir(parents=True, exist_ok=True)
        
        print(f"🎯 CacBuocThucHienProcessor initialized - step counting logic")

    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"❌ Error reading {pdf_path}: {str(e)}")
            return None

    def count_steps_in_chuong_v(self, chuong_v_content):
        """Count the number of steps in the process table from CHUONG_V.pdf"""
        prompt = f"""
From the CHUONG V document below, find the section about "Yêu cầu về quy trình chỉnh lý" or similar wording.

Look for a table that lists implementation steps. This table typically has:
- Step numbers (like 1, 2, 3, etc.)
- Process descriptions for each step
- The steps are usually numbered sequentially

Your task:
1. Find this process steps table
2. Count the TOTAL number of steps listed
3. The count should be either 21 or 23 steps

Return ONLY the number (21 or 23). If you can't find the table or the count is different, return "UNKNOWN".

CHUONG V CONTENT:
{chuong_v_content}

STEP COUNT:"""

        try:
            response = openai.ChatCompletion.create(
                model='gpt-4o',
                messages=[
                    {"role": "system", "content": "You are an expert at analyzing Vietnamese document structures and counting process steps in tables. Return only the step count number."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=50,
                temperature=0.0
            )
            
            step_count = response.choices[0].message.content.strip()
            print(f"🔢 OpenAI detected step count: '{step_count}'")
            
            # Validate the response
            if step_count in ["21", "23"]:
                return int(step_count)
            else:
                print(f"⚠️ Unexpected step count: {step_count}")
                return None
            
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return None

    def select_source_file(self, step_count):
        """Select the appropriate source file based on step count"""
        if step_count == 21:
            source_file = "21_BUOC.docx"
        elif step_count == 23:
            source_file = "23_BUOC.docx"
        else:
            raise ValueError(f"Invalid step count: {step_count}. Expected 21 or 23.")
        
        source_path = Path(source_file)
        if not source_path.exists():
            raise FileNotFoundError(f"Source file not found: {source_file}")
        
        print(f"📄 Selected source file: {source_file}")
        return source_path

    def copy_source_to_processed(self, source_path):
        """Copy the selected source file to processed folder as output.docx"""
        dest_path = self.process_folder / 'output.docx'
        shutil.copy2(source_path, dest_path)
        print(f"✅ Copied {source_path.name} to {dest_path}")
        return dest_path

    def replace_placeholder(self, doc, placeholder_tag):
        """Replace placeholder using proven method"""
        pattern = r"\{\{(.+?)\}\}"
        match = re.search(pattern, placeholder_tag)
        if not match:
            raise ValueError(f"Invalid placeholder: {placeholder_tag}")
        
        tag_name = match.group(1)
        source_path = Path(f"processed/{tag_name}/output.docx")

        print(f"🔍 Looking for source file: {source_path}")
        if not source_path.exists():
            raise FileNotFoundError(f"Missing source doc: {source_path}")

        source_doc = Document(source_path)
        
        # Debug: Check what's in the source document
        print(f"📊 Source document has {len(source_doc.paragraphs)} paragraphs")
        print(f"📊 Source document has {len(source_doc.tables)} tables")
        
        # Get all content (paragraphs and tables)
        all_elements = []
        
        # Add paragraphs
        for para in source_doc.paragraphs:
            if para.text.strip():
                all_elements.append(('paragraph', para))
        
        # Add tables
        for table in source_doc.tables:
            all_elements.append(('table', table))
        
        print(f"📊 Total elements to copy: {len(all_elements)}")

        # Find and replace placeholder
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                print(f"📍 Found placeholder in paragraph {i}")
                
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)
                
                print(f"🔄 Copying {len(all_elements)} elements...")

                # Insert all elements from source
                for element_type, element in reversed(all_elements):
                    if element_type == 'paragraph':
                        new_p = deepcopy(element._element)
                        parent.insert(index, new_p)
                    elif element_type == 'table':
                        new_t = deepcopy(element._element)
                        parent.insert(index, new_t)
                
                print(f"✅ Successfully copied all elements")
                break

    def copy_template_to_output(self):
        """Copy template file to output file"""
        try:
            if not os.path.exists(self.template_file):
                print(f"❌ Template file not found: {self.template_file}")
                return False
            
            shutil.copy2(self.template_file, self.output_file)
            print(f"✅ Copied template to output file: {self.output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error copying template: {str(e)}")
            return False

    def test_cac_buoc_thuc_hien_process(self):
        """Test the complete step counting and replacement process"""
        print("\n🧪 TESTING: {{cac_buoc_thuc_hien}} - Step Counting Logic")
        print("=" * 60)
        
        # Check if CHUONG_V.pdf exists
        chuong_v_file = self.pdf_folder / 'CHUONG_V.pdf'
        if not chuong_v_file.exists():
            print(f"❌ File not found: {chuong_v_file}")
            return False
        
        # Step 1: Extract text from CHUONG_V.pdf
        print("📖 Step 1: Reading CHUONG_V.pdf...")
        chuong_v_content = self.extract_text_from_pdf(chuong_v_file)
        if not chuong_v_content:
            return False
        
        print(f"✅ Extracted {len(chuong_v_content)} characters")
        
        # Step 2: Count steps in the process table
        print("🔢 Step 2: Counting process steps...")
        step_count = self.count_steps_in_chuong_v(chuong_v_content)
        
        if not step_count:
            print("❌ Failed to determine step count")
            return False
        
        print(f"✅ Detected {step_count} steps")
        
        # Step 3: Select appropriate source file
        print(f"📄 Step 3: Selecting source file for {step_count} steps...")
        try:
            source_path = self.select_source_file(step_count)
        except (ValueError, FileNotFoundError) as e:
            print(f"❌ {e}")
            return False
        
        # Step 4: Copy source to processed folder
        print("🔄 Step 4: Preparing source document...")
        processed_path = self.copy_source_to_processed(source_path)
        
        # Step 5: Replace placeholder in template
        print("🔄 Step 5: Replacing placeholder in template...")
        if self.copy_template_to_output():
            doc = Document(self.output_file)
            self.replace_placeholder(doc, "{{cac_buoc_thuc_hien}}")
            doc.save(self.output_file)
            print(f"✅ SUCCESS: {{cac_buoc_thuc_hien}} processed with {step_count} steps!")
            print(f"📄 Check output file: {self.output_file}")
            return True
        
        return False

def main():
    print("🇻🇳 Cac Buoc Thuc Hien Processor - Step Counting Logic")
    print("=" * 60)
    
    try:
        processor = CacBuocThucHienProcessor()
        print("✅ Environment loaded successfully")
        
        processor.test_cac_buoc_thuc_hien_process()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()