import os
import shutil
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from dotenv import load_dotenv
import re
from copy import deepcopy

# Load environment variables
load_dotenv()

class CacBuocTextProcessor:
    def __init__(self):
        """Initialize the processor for cac_buoc_text - find location then extract with PyMuPDF"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        
        self.pdf_folder = Path("pdf_inputs")
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        
        # Create processing folder structure
        self.process_folder = Path("processed/cac_buoc_text")
        self.process_folder.mkdir(parents=True, exist_ok=True)
        
        print("🎯 CacBuocTextProcessor initialized - AI LOCATION + PYMUPDF EXTRACTION")

    def extract_text_from_pdf_precise(self, pdf_path):
        """Extract text with precise positioning using PyMuPDF"""
        try:
            doc = fitz.open(pdf_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                full_text += f"\n--- PAGE {page_num + 1} ---\n"
                full_text += text
            
            doc.close()
            return full_text
        except Exception as e:
            print(f"❌ Error reading PDF with PyMuPDF: {str(e)}")
            return None

    def find_text_boundaries(self, pdf_content):
        """Use AI to find the boundaries - heading and table locations"""
        prompt = f"""
TASK: Find the text boundaries in this PDF content.

STEP 1: Find a table that contains "Nội dung công việc" (or similar wording like "Số TT")
STEP 2: Working backwards from that table, find the nearest heading before it (like 3.1, 3.2, 3.3, etc.)

I need to extract the text that appears:
- AFTER the heading (like "3.3 Yêu cầu về quy trình chỉnh lý")  
- BEFORE the table (that has "Số TT" and "Nội dung công việc")

Please identify these boundaries by finding specific text phrases I can use to locate them.

Return format:
HEADING_TEXT: [the exact heading text to search for]
TABLE_START_TEXT: [the exact text that marks where the table begins]

PDF CONTENT:
{pdf_content}

BOUNDARIES:
"""

        try:
            response = openai.ChatCompletion.create(
                model='gpt-4o',
                messages=[
                    {"role": "system", "content": "Find the text boundaries by identifying the heading before and table after the target text. Return exact text phrases for location."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.0
            )
            
            result = response.choices[0].message.content.strip()
            print("🎯 AI found text boundaries")
            
            # Save for debugging
            (self.process_folder / 'boundaries.txt').write_text(result, encoding='utf-8')
            
            return result
            
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return None

    def extract_text_between_boundaries(self, pdf_content, boundaries):
        """Use PyMuPDF to extract exact text between boundaries - FLEXIBLE MATCHING"""
        lines = boundaries.split('\n')
        
        heading_text = ""
        table_start_text = ""
        
        for line in lines:
            line = line.strip()
            if line.startswith('HEADING_TEXT:'):
                heading_text = line.replace('HEADING_TEXT:', '').strip().strip('"')
            elif line.startswith('TABLE_START_TEXT:'):
                table_start_text = line.replace('TABLE_START_TEXT:', '').strip().strip('"')
        
        print(f"🔍 AI suggested boundaries:")
        print(f"  Heading: '{heading_text}'")
        print(f"  Table: '{table_start_text}'")
        
        # FLEXIBLE HEADING SEARCH - look for key parts
        heading_pos = -1
        if heading_text:
            # Try exact match first (without quotes)
            heading_pos = pdf_content.find(heading_text)
            
            # If not found, try flexible search for key words
            if heading_pos == -1:
                # Extract key words from heading
                if "quy trình chỉnh lý" in heading_text.lower():
                    heading_pos = pdf_content.lower().find("quy trình chỉnh lý")
                elif "công việc thực hiện" in heading_text.lower():
                    heading_pos = pdf_content.lower().find("công việc thực hiện")
                elif "yêu cầu" in heading_text.lower():
                    heading_pos = pdf_content.lower().find("yêu cầu")
        
        # FLEXIBLE TABLE SEARCH - try multiple patterns
        table_pos = -1
        if table_start_text:
            # Try exact match first
            table_pos = pdf_content.find(table_start_text)
            
            # If not found, try flexible search with different patterns
            if table_pos == -1:
                search_patterns = [
                    "số tt",
                    "nội dung công việc", 
                    "ghi chú",
                    "số\ntt",  # Multi-line pattern
                    "nội dung\ncông việc",  # Multi-line pattern
                ]
                
                print("🔍 Trying table search patterns:")
                for pattern in search_patterns:
                    pos = pdf_content.lower().find(pattern)
                    print(f"  '{pattern}' -> position: {pos}")
                    if pos != -1:
                        table_pos = pos
                        print(f"  ✅ Found table using pattern: '{pattern}'")
                        break
        
        print(f"🔍 Found positions:")
        print(f"  Heading at: {heading_pos}")
        print(f"  Table at: {table_pos}")
        
        if heading_pos == -1:
            print(f"❌ Could not find heading with flexible search")
            return ""
        
        if table_pos == -1:
            print(f"❌ Could not find table with flexible search")
            return ""
        
        if table_pos <= heading_pos:
            print("❌ Table appears before heading - check boundaries")
            return ""
        
        # Find the end of the heading line
        heading_line_end = pdf_content.find('\n', heading_pos)
        if heading_line_end == -1:
            heading_line_end = heading_pos + len(heading_text)
        
        # Extract text between end of heading line and start of table
        start_pos = heading_line_end
        extracted_text = pdf_content[start_pos:table_pos].strip()
        
        # Clean up: remove "Số TT" variations that might appear at the end
        print(f"🔍 Text ending (last 20 chars): '{extracted_text[-20:]}'")
        print(f"🔍 Text ending repr: {repr(extracted_text[-20:])}")
        
        cleanup_patterns = [
            "Số TT",           # Same line
            ". Số TT",         # After period same line
            "Số\nTT",          # Multi-line
            ".\nSố\nTT",       # After period multi-line
            "\nSố\nTT",        # Multi-line with leading newline
            " Số\nTT",         # With space
            ".\n\nSố\nTT",     # Double newline
        ]
        
        original_length = len(extracted_text)
        
        for pattern in cleanup_patterns:
            if extracted_text.endswith(pattern):
                extracted_text = extracted_text[:-len(pattern)].strip()
                print(f"🧹 Cleaned up pattern: '{pattern.replace(chr(10), '\\n')}'")
                print(f"🧹 Removed {original_length - len(extracted_text)} characters")
                break
        else:
            print("⚠️  No cleanup pattern matched - checking manual removal")
            # Manual cleanup if patterns don't work
            if "Số" in extracted_text[-10:] and "TT" in extracted_text[-10:]:
                # Find last occurrence of "dự án" and cut there
                last_du_an = extracted_text.rfind("dự án")
                if last_du_an != -1:
                    # Find the end of this phrase (usually followed by period)
                    end_pos = last_du_an + 5  # "dự án" is 5 chars
                    if end_pos < len(extracted_text) and extracted_text[end_pos] == '.':
                        end_pos += 1
                    extracted_text = extracted_text[:end_pos].strip()
                    print(f"🧹 Manual cleanup: cut after 'dự án.' - removed {original_length - len(extracted_text)} characters")
        
        print(f"✅ Extracted {len(extracted_text)} characters between boundaries")
        print(f"📄 First 200 chars: '{extracted_text[:200]}...'")
        print(f"📄 Last 50 chars: '...{extracted_text[-50:]}'")
        
        # Save for debugging
        (self.process_folder / 'extracted_text.txt').write_text(extracted_text, encoding='utf-8')
        
        return extracted_text

    def parse_text_into_paragraphs(self, extracted_text):
        """Parse the extracted text into paragraphs for DOCX formatting"""
        if not extracted_text:
            return []
        
        # Split by lines that start with dash or by double line breaks
        lines = extracted_text.split('\n')
        paragraphs = []
        current_paragraph = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                # Empty line - end current paragraph if it exists
                if current_paragraph:
                    paragraphs.append(current_paragraph.strip())
                    current_paragraph = ""
            elif line.startswith('-'):
                # New paragraph starting with dash
                if current_paragraph:
                    paragraphs.append(current_paragraph.strip())
                current_paragraph = line
            else:
                # Continuation of current paragraph
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
        
        # Add the last paragraph if it exists
        if current_paragraph:
            paragraphs.append(current_paragraph.strip())
        
        print(f"📄 Parsed into {len(paragraphs)} paragraphs:")
        for i, para in enumerate(paragraphs):
            print(f"  Para{i+1}: {len(para)} chars - '{para[:100]}...'")
        
        return paragraphs

    def create_text_docx(self, paragraphs):
        """Create DOCX with properly formatted paragraphs"""
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # Add all paragraphs with proper formatting
        for i, paragraph in enumerate(paragraphs):
            if paragraph:
                para = doc.add_paragraph(paragraph)
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = 1.15
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    
                print(f"✅ Added paragraph {i+1}")

        # Save
        docx_path = self.process_folder / 'output.docx'
        doc.save(docx_path)
        print(f"✅ Created text DOCX with {len(paragraphs)} paragraphs: {docx_path}")
        
        return docx_path

    def replace_placeholder(self, doc, placeholder_tag):
        """Replace placeholder using proven method"""
        pattern = r"\{\{(.+?)\}\}"
        match = re.search(pattern, placeholder_tag)
        if not match:
            raise ValueError(f"Invalid placeholder: {placeholder_tag}")
        
        tag_name = match.group(1)
        source_path = Path(f"processed/{tag_name}/output.docx")

        if not source_path.exists():
            raise FileNotFoundError(f"Missing source doc: {source_path}")

        source_doc = Document(source_path)
        
        # Get all content
        all_elements = []
        for para in source_doc.paragraphs:
            if para.text.strip():
                all_elements.append(('paragraph', para))
        for table in source_doc.tables:
            all_elements.append(('table', table))

        # Replace placeholder
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)

                for element_type, element in reversed(all_elements):
                    if element_type == 'paragraph':
                        new_p = deepcopy(element._element)
                        parent.insert(index, new_p)
                    elif element_type == 'table':
                        new_t = deepcopy(element._element)
                        parent.insert(index, new_t)
                break

    def copy_template_to_output(self):
        """Copy template file to output file"""
        try:
            if not os.path.exists(self.template_file):
                print(f"❌ Template file not found: {self.template_file}")
                return False
            
            shutil.copy2(self.template_file, self.output_file)
            print(f"✅ Copied template to output file: {self.output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error copying template: {str(e)}")
            return False

    def test_text_extraction(self):
        """Test AI location + PyMuPDF extraction"""
        print("\n🧪 TESTING: {{cac_buoc_text}} - AI LOCATION + PYMUPDF EXTRACTION")
        print("=" * 60)
        
        # Check if CHUONG_V.pdf exists
        chuong_v_file = self.pdf_folder / 'CHUONG_V.pdf'
        if not chuong_v_file.exists():
            print(f"❌ File not found: {chuong_v_file}")
            return False
        
        # Step 1: Extract with PyMuPDF
        print("📖 Step 1: Reading PDF with PyMuPDF...")
        pdf_content = self.extract_text_from_pdf_precise(chuong_v_file)
        if not pdf_content:
            return False
        
        print(f"✅ Extracted {len(pdf_content)} characters")
        
        # Step 2: Use AI to find boundaries
        print("🔍 Step 2: Using AI to find text boundaries...")
        boundaries = self.find_text_boundaries(pdf_content)
        if not boundaries:
            return False
        
        # Step 3: Use PyMuPDF to extract exact text between boundaries
        print("📋 Step 3: Using PyMuPDF to extract exact text...")
        extracted_text = self.extract_text_between_boundaries(pdf_content, boundaries)
        if not extracted_text:
            return False
        
        # Step 4: Parse into paragraphs
        print("📄 Step 4: Parsing text into paragraphs...")
        paragraphs = self.parse_text_into_paragraphs(extracted_text)
        
        # Step 5: Create DOCX
        print("📝 Step 5: Creating formatted DOCX...")
        try:
            docx_path = self.create_text_docx(paragraphs)
        except Exception as e:
            print(f"❌ Failed to create DOCX: {str(e)}")
            return False
        
        # Step 6: Replace placeholder
        print("🔄 Step 6: Replacing placeholder...")
        if self.copy_template_to_output():
            doc = Document(self.output_file)
            self.replace_placeholder(doc, "{{cac_buoc_text}}")
            doc.save(self.output_file)
            print(f"✅ SUCCESS: AI location + PyMuPDF extraction completed!")
            print(f"📄 Check output file: {self.output_file}")
            return True
        
        return False

def main():
    print("🇻🇳 Cac Buoc Text Processor - AI LOCATION + PYMUPDF EXTRACTION")
    print("=" * 60)
    
    try:
        processor = CacBuocTextProcessor()
        print("✅ Environment loaded successfully")
        
        processor.test_text_extraction()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()