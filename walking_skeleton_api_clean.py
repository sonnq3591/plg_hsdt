#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Walking Skeleton FastAPI - Clean Production Version
Wraps your proven modular workflow in a REST API
"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import shutil
import tempfile
from pathlib import Path
from typing import List
import uvicorn
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Vietnamese Procurement Document API",
    description="Walking Skeleton: Process 02_MUC_DO_HIEU_BIET template with all 5 placeholders",
    version="1.0.0"
)

# Add CORS middleware for web testing
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class WalkingSkeletonProcessor:
    """
    Production version of your proven walking skeleton workflow
    Reuses all your existing modules exactly as they work locally
    """
    
    def __init__(self):
        self.work_dir = None
        self.pdf_inputs_dir = None
        self.required_pdfs = ['TBMT.pdf', 'BMMT.pdf', 'CHUONG_III.pdf', 'CHUONG_V.pdf', 'HSMT.pdf']
        
    def setup_workspace(self):
        """Create temporary workspace for processing"""
        self.work_dir = Path(tempfile.mkdtemp(prefix="walking_skeleton_"))
        self.pdf_inputs_dir = self.work_dir / "pdf_inputs"
        self.pdf_inputs_dir.mkdir(exist_ok=True)
        logger.info(f"Created workspace: {self.work_dir}")
        
    def save_uploaded_pdfs(self, pdf_files: List[UploadFile]):
        """Save uploaded PDFs to workspace"""
        saved_files = {}
        
        for pdf_file in pdf_files:
            if pdf_file.filename in self.required_pdfs:
                file_path = self.pdf_inputs_dir / pdf_file.filename
                with open(file_path, "wb") as f:
                    shutil.copyfileobj(pdf_file.file, f)
                saved_files[pdf_file.filename] = file_path
                logger.info(f"Saved: {pdf_file.filename} ({file_path.stat().st_size} bytes)")
            else:
                logger.warning(f"Unexpected file: {pdf_file.filename}")
        
        return saved_files
    
    def copy_required_templates(self):
        """Copy required template files to workspace"""
        required_files = [
            "02_MUC_DO_HIEU_BIET_template.docx",
            "21_BUOC.docx",
            "23_BUOC.docx"
        ]
        
        for file_name in required_files:
            source = Path(file_name)
            if source.exists():
                dest = self.work_dir / file_name
                shutil.copy2(source, dest)
                logger.info(f"Copied template: {file_name}")
            else:
                raise FileNotFoundError(f"Required template not found: {file_name}")
    
    def process_all_placeholders(self):
        """
        Execute your proven walking skeleton workflow
        All 5 placeholders processed sequentially with accumulation
        """
        logger.info("Starting walking skeleton workflow...")
        
        # Change to workspace directory
        original_cwd = os.getcwd()
        os.chdir(self.work_dir)
        
        try:
            # Step 1: {{ten_goi_thau}} - starts the chain
            logger.info("Step 1/5: Processing {{ten_goi_thau}}...")
            from processor import VietnameseProcurementProcessor
            processor1 = VietnameseProcurementProcessor()
            success1 = processor1.test_ten_goi_thau_extraction()
            if not success1:
                raise Exception("Failed Step 1: {{ten_goi_thau}}")
            logger.info("✅ Step 1 complete: {{ten_goi_thau}}")
            
            # Step 2: {{pham_vi_cung_cap}} - accumulates on Step 1 result
            logger.info("Step 2/5: Processing {{pham_vi_cung_cap}}...")
            from processor_pham_vi import PhamViCungCapProcessor
            processor2 = PhamViCungCapProcessor()
            processor2.template_file = "02_MUC_DO_HIEU_BIET_output.docx"
            processor2.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
            
            # Smart copy override
            original_copy = processor2.copy_template_to_output
            def smart_copy():
                if processor2.template_file == processor2.output_file:
                    logger.info("Same file - no copy needed")
                    return True
                return original_copy()
            processor2.copy_template_to_output = smart_copy
            
            success2 = processor2.test_pham_vi_cung_cap_simple()
            if not success2:
                raise Exception("Failed Step 2: {{pham_vi_cung_cap}}")
            logger.info("✅ Step 2 complete: {{pham_vi_cung_cap}}")
            
            # Step 3: {{can_cu_phap_ly}} - accumulates on Step 2 result
            logger.info("Step 3/5: Processing {{can_cu_phap_ly}}...")
            from processor_can_cu import CanCuPhapLyProcessor
            processor3 = CanCuPhapLyProcessor()
            processor3.template_file = "02_MUC_DO_HIEU_BIET_output.docx"
            processor3.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
            
            # Smart copy override
            original_copy3 = processor3.copy_template_to_output
            def smart_copy3():
                if processor3.template_file == processor3.output_file:
                    logger.info("Same file - no copy needed")
                    return True
                return original_copy3()
            processor3.copy_template_to_output = smart_copy3
            
            success3 = processor3.test_can_cu_phap_ly_full_process()
            if not success3:
                raise Exception("Failed Step 3: {{can_cu_phap_ly}}")
            logger.info("✅ Step 3 complete: {{can_cu_phap_ly}}")
            
            # Step 4: {{muc_dich_cong_viec}} - accumulates on Step 3 result
            logger.info("Step 4/5: Processing {{muc_dich_cong_viec}}...")
            from processor_muc_dich import MucDichProcessor
            processor4 = MucDichProcessor()
            processor4.template_file = "02_MUC_DO_HIEU_BIET_output.docx"
            processor4.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
            
            # Smart copy override
            original_copy4 = processor4.copy_template_to_output
            def smart_copy4():
                if processor4.template_file == processor4.output_file:
                    logger.info("Same file - no copy needed")
                    return True
                return original_copy4()
            processor4.copy_template_to_output = smart_copy4
            
            success4 = processor4.test_muc_dich_cong_viec_full_process()
            if not success4:
                raise Exception("Failed Step 4: {{muc_dich_cong_viec}}")
            logger.info("✅ Step 4 complete: {{muc_dich_cong_viec}}")
            
            # Step 5: {{cac_buoc_thuc_hien}} - final accumulation
            logger.info("Step 5/5: Processing {{cac_buoc_thuc_hien}}...")
            from combined_processor import CombinedProcessor
            
            processor5 = CombinedProcessor()
            processor5.template_file = "02_MUC_DO_HIEU_BIET_output.docx"
            processor5.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
            
            # FIX: Update the PDF path for API workspace
            processor5.chuong_v_pdf = "pdf_inputs/CHUONG_V.pdf"
            
            # Debug: Check if files exist
            template_exists = Path("02_MUC_DO_HIEU_BIET_output.docx").exists()
            buoc21_exists = Path("21_BUOC.docx").exists()
            buoc23_exists = Path("23_BUOC.docx").exists()
            chuong_v_exists = Path("pdf_inputs/CHUONG_V.pdf").exists()
            logger.info(f"Files check - Template: {template_exists}, 21_BUOC: {buoc21_exists}, 23_BUOC: {buoc23_exists}, CHUONG_V: {chuong_v_exists}")
            
            # Override the replace method for same file handling
            original_replace = processor5.replace_placeholder_only
            def smart_replace_final(template_path, content_path, output_path, placeholder="{{cac_buoc_thuc_hien}}"):
                logger.info(f"Smart replace: {template_path} -> {output_path}")
                
                if template_path == output_path:
                    logger.info("Working directly on file")
                    
                    from docx import Document
                    from copy import deepcopy
                    
                    doc = Document(output_path)
                    source_doc = Document(content_path)
                    
                    # Get content elements
                    all_elements = []
                    for para in source_doc.paragraphs:
                        if para.text.strip():
                            all_elements.append(('paragraph', para))
                    for table in source_doc.tables:
                        all_elements.append(('table', table))
                    
                    logger.info(f"Found {len(all_elements)} elements to copy")
                    
                    # Find and replace placeholder
                    for i, paragraph in enumerate(doc.paragraphs):
                        if placeholder in paragraph.text:
                            logger.info(f"Found placeholder in paragraph {i}")
                            
                            p_element = paragraph._element
                            parent = p_element.getparent()
                            index = parent.index(p_element)
                            parent.remove(p_element)
                            
                            # Insert all elements
                            for element_type, element in reversed(all_elements):
                                if element_type == 'paragraph':
                                    new_p = deepcopy(element._element)
                                    parent.insert(index, new_p)
                                elif element_type == 'table':
                                    new_t = deepcopy(element._element)
                                    parent.insert(index, new_t)
                            break
                    
                    doc.save(output_path)
                    logger.info("Document saved successfully")
                    return True
                else:
                    return original_replace(template_path, content_path, output_path, placeholder)
            
            processor5.replace_placeholder_only = smart_replace_final
            
            # Execute Step 5
            success5 = processor5.process_complete_workflow()
            if not success5:
                raise Exception("Failed Step 5: {{cac_buoc_thuc_hien}}")
            
            logger.info("✅ Step 5 complete: {{cac_buoc_thuc_hien}}")
            
            # Verify final output
            output_file = self.work_dir / "02_MUC_DO_HIEU_BIET_output.docx"
            if not output_file.exists():
                raise Exception("Final output file not generated")
            
            file_size = output_file.stat().st_size
            logger.info(f"🎉 Walking skeleton complete! File size: {file_size:,} bytes")
            return output_file
            
        finally:
            os.chdir(original_cwd)
    
    def cleanup(self):
        """Clean up workspace"""
        if self.work_dir and self.work_dir.exists():
            shutil.rmtree(self.work_dir)
            logger.info("Workspace cleaned up")

# API Endpoints

@app.get("/")
async def root():
    """API root - basic info"""
    return {
        "message": "Vietnamese Procurement Document API - Walking Skeleton",
        "status": "running",
        "template": "02_MUC_DO_HIEU_BIET",
        "placeholders": 5,
        "endpoints": {
            "process": "/api/process-document",
            "health": "/api/health",
            "docs": "/docs"
        }
    }

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "message": "Walking Skeleton API operational",
        "version": "1.0.0"
    }

@app.post("/api/process-document")
async def process_document(
    tbmt_pdf: UploadFile = File(..., description="TBMT.pdf - Thông báo mời thầu"),
    bmmt_pdf: UploadFile = File(..., description="BMMT.pdf - Biểu mẫu mời thầu"),
    chuong_iii_pdf: UploadFile = File(..., description="CHUONG_III.pdf"),
    chuong_v_pdf: UploadFile = File(..., description="CHUONG_V.pdf"),
    hsmt_pdf: UploadFile = File(..., description="HSMT.pdf - Hồ sơ mời thầu")
):
    """
    Walking Skeleton Endpoint: Process 02_MUC_DO_HIEU_BIET template
    
    Uploads 5 PDFs → Processes all 5 placeholders → Returns complete DOCX
    
    Returns: 02_MUC_DO_HIEU_BIET_output.docx with all placeholders replaced
    """
    processor = WalkingSkeletonProcessor()
    
    try:
        logger.info("=== Walking Skeleton API Request Started ===")
        
        # Setup workspace
        processor.setup_workspace()
        logger.info("Workspace created")
        
        # Copy templates
        processor.copy_required_templates()
        logger.info("Templates copied")
        
        # Save uploaded PDFs
        pdf_files = [tbmt_pdf, bmmt_pdf, chuong_iii_pdf, chuong_v_pdf, hsmt_pdf]
        saved_files = processor.save_uploaded_pdfs(pdf_files)
        logger.info(f"PDFs saved: {list(saved_files.keys())}")
        
        # Validate all required PDFs
        required = set(processor.required_pdfs)
        uploaded = set(saved_files.keys())
        missing = required - uploaded
        
        if missing:
            raise HTTPException(
                status_code=400, 
                detail=f"Missing required PDF files: {list(missing)}"
            )
        
        # Process all placeholders using proven workflow
        output_file = processor.process_all_placeholders()
        logger.info("All placeholders processed successfully")
        
        # Copy the file to a safe location before cleanup
        safe_output = Path(tempfile.gettempdir()) / "api_result.docx"
        shutil.copy2(output_file, safe_output)
        logger.info(f"File copied to safe location: {safe_output}")
        
        # Clean up workspace
        processor.cleanup()
        
        # Return the generated DOCX from safe location
        return FileResponse(
            path=safe_output,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename="02_MUC_DO_HIEU_BIET_output.docx",
            headers={"Content-Disposition": "attachment; filename=02_MUC_DO_HIEU_BIET_output.docx"}
        )
        
    except Exception as e:
        logger.error(f"Processing failed: {str(e)}")
        processor.cleanup()  # Clean up on error
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@app.get("/api/templates")
async def list_templates():
    """List available templates and their status"""
    return {
        "walking_skeleton": {
            "template": "02_MUC_DO_HIEU_BIET",
            "status": "✅ IMPLEMENTED",
            "placeholders": [
                "{{ten_goi_thau}}",
                "{{pham_vi_cung_cap}}",
                "{{can_cu_phap_ly}}",
                "{{muc_dich_cong_viec}}",
                "{{cac_buoc_thuc_hien}}"
            ],
            "required_pdfs": [
                "TBMT.pdf",
                "BMMT.pdf", 
                "CHUONG_III.pdf",
                "CHUONG_V.pdf",
                "HSMT.pdf"
            ]
        },
        "future_templates": {
            "count": 14,
            "status": "🔄 TODO",
            "note": "Will be added incrementally after Teams bot integration"
        }
    }

if __name__ == "__main__":
    print("🚀 Starting Walking Skeleton API...")
    print("🎯 Template: 02_MUC_DO_HIEU_BIET.docx")
    print("📋 Placeholders: 5 total")
    print("🌐 Server: http://localhost:8000")
    print("📖 API docs: http://localhost:8000/docs")
    print("✅ Ready for Teams bot integration!")
    
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
        log_level="info"
    )