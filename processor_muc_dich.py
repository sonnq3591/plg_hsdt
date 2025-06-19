import os
import shutil
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from dotenv import load_dotenv
import re
from copy import deepcopy

# Load environment variables
load_dotenv()

class MucDichProcessor:
    def __init__(self):
        """Initialize the processor following your proven full process"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        
        self.pdf_folder = Path("pdf_inputs")
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        
        # Create processing folder structure like your approach
        self.process_folder = Path("processed/muc_dich_cong_viec")
        self.process_folder.mkdir(parents=True, exist_ok=True)
        
        print(f"🎯 MucDichProcessor initialized using full proven process")

    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"❌ Error reading {pdf_path}: {str(e)}")
            return None

    def format_text_markdown(self, input_text, system_prompt):
        """Your proven OpenAI formatting function"""
        try:
            response = openai.ChatCompletion.create(
                model='gpt-4o',
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": input_text}
                ],
                max_tokens=1500,
                temperature=0.0
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return "[KHÔNG TÌM THẤY]"

    def create_system_prompt(self):
        """Create system prompt that excludes legal basis sections"""
        return """You are an assistant that extracts structured content related to "Mục đích" or "Mục đích Công việc" or similar wording from Vietnamese construction documents.

Your task is to:
- Intelligently identify the section that presents the content for "Mục đích" or "Mục đích Công việc" or similar wording, even if the section title is phrased differently.
- Extract its full content **without** modifying or summarizing anything.
- SKIP any section titles, headers, or sub-headers like "Mục đích công việc:", "Mục tiêu công việc:", "a) Mục đích:", etc.
- Start directly with the actual content paragraphs.
- **STOP extraction when you encounter legal basis sections** like "Căn cứ pháp lý", "Căn cứ Luật", "Quy định về", "Các Văn bản Luật", etc.

FORMAT STRICTLY as Markdown, following these rules:

1. **DO NOT include any leading section prefix** like 'c)', 'a.', '1)', etc.
2. **DO NOT include section headers** like "Mục đích công việc:", "Mục tiêu công việc:", etc.
3. **DO NOT include legal basis content** - stop before any "Căn cứ pháp lý" or similar sections.
4. The very first line (intro/statement) must NOT be bolded.
5. Only bold subheadings that are **followed by a list** (e.g. lines starting with '-').
6. Preserve all line breaks and paragraph groupings as they are in the original.
7. Every line (intro, subheading, and bullet) must be indented with **0.5 tab** (2 spaces).
8. Use dash '-' for bullet points and maintain their original text.
9. Do NOT paraphrase, summarize, reword, remove, or insert any content. Keep the original words exactly.
10. Extract ONLY the purpose/objective content, nothing about legal references or regulations."""

    def process_to_markdown(self, input_text):
        """Step 1: Extract and format to markdown"""
        system_prompt = self.create_system_prompt()
        markdown_output = self.format_text_markdown(input_text, system_prompt)
        
        # Save markdown like your approach
        (self.process_folder / 'output.md').write_text(markdown_output, encoding='utf-8')
        print(f"✅ Saved markdown: {self.process_folder}/output.md")
        
        return markdown_output

    def markdown_to_docx(self, markdown_content):
        """Step 2: Convert markdown to DOCX using your proven method"""
        # Create DOCX using your exact approach
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        for line in markdown_content.splitlines():
            line = line.strip()
            if not line:
                continue
            if re.match(r'^\*\*(.+)\*\*$', line):
                text = re.sub(r'\*\*(.+)\*\*', r'\1', line)
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
            elif line.startswith('- '):
                para = doc.add_paragraph(f'- {line[2:].strip()}')
            else:
                para = doc.add_paragraph(line)
            
            # Your proven formatting
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)

        # Save DOCX like your approach
        docx_path = self.process_folder / 'output.docx'
        doc.save(docx_path)
        print(f"✅ Saved DOCX: {docx_path}")
        
        return docx_path

    def replace_text_variables_preserve_runs(self, doc, var_dict):
        """Your proven cross-run replacement method"""
        for paragraph in doc.paragraphs:
            i = 0
            while i < len(paragraph.runs):
                # Try to match across runs
                run_text = ""
                j = i
                while j < len(paragraph.runs) and len(run_text) < 100:
                    run_text += paragraph.runs[j].text
                    j += 1

                    for key, val in var_dict.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in run_text:
                            # Split into 3 parts: before, replacement, after
                            before, after = run_text.split(placeholder, 1)

                            # Clear affected runs
                            for k in range(i, j):
                                paragraph.runs[k].text = ""

                            # Write back: before + replacement + after using original styles
                            if before:
                                paragraph.runs[i].text = before
                            paragraph.runs[i + 1].text = val
                            if after:
                                paragraph.runs[i + 2].text = after

                            i = j  # move past replaced section
                            break
                    else:
                        continue  # inner loop didn't break
                    break  # outer loop: matched a placeholder, break

                i += 1

    def replace_placeholder(self, doc, placeholder_tag):
        """Your proven placeholder replacement method"""
        # Extract folder name from tag, e.g. {{muc_dich_cong_viec}} -> muc_dich_cong_viec
        match = re.search(r"\{\{(.+?)\}\}", placeholder_tag)
        if not match:
            raise ValueError(f"Invalid placeholder: {placeholder_tag}")
        
        tag_name = match.group(1)
        source_path = Path(f"processed/{tag_name}/output.docx")

        if not source_path.exists():
            raise FileNotFoundError(f"Missing source doc: {source_path}")

        source_doc = Document(source_path)
        source_paragraphs = [p for p in source_doc.paragraphs if p.text.strip()]

        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)

                for src_p in reversed(source_paragraphs):
                    new_p = deepcopy(src_p._element)
                    parent.insert(index, new_p)
                    inserted_p = doc.paragraphs[index]

                    inserted_p.paragraph_format.line_spacing = src_p.paragraph_format.line_spacing or 1.3
                    inserted_p.paragraph_format.space_before = src_p.paragraph_format.space_before
                    inserted_p.paragraph_format.space_after = src_p.paragraph_format.space_after
                    inserted_p.paragraph_format.left_indent = src_p.paragraph_format.left_indent
                    inserted_p.paragraph_format.first_line_indent = src_p.paragraph_format.first_line_indent

                    for run_idx, run in enumerate(inserted_p.runs):
                        try:
                            src_run = src_p.runs[run_idx]
                            run.font.size = src_run.font.size or Pt(14)
                            run.font.name = src_run.font.name or "Times New Roman"
                        except IndexError:
                            run.font.size = Pt(14)
                            run.font.name = "Times New Roman"
                break

    def copy_template_to_output(self):
        """Copy template file to output file"""
        try:
            if not os.path.exists(self.template_file):
                print(f"❌ Template file not found: {self.template_file}")
                return False
            
            shutil.copy2(self.template_file, self.output_file)
            print(f"✅ Copied template to output file: {self.output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error copying template: {str(e)}")
            return False

    def test_muc_dich_cong_viec_full_process(self):
        """Test using your complete proven process"""
        print("\n🧪 TESTING: {{muc_dich_cong_viec}} - Full Proven Process")
        print("=" * 60)
        
        # Check if CHUONG_V.pdf exists
        chuong_v_file = self.pdf_folder / self.pdf_files['CHUONG_V']
        if not chuong_v_file.exists():
            print(f"❌ File not found: {chuong_v_file}")
            return False
        
        # Step 1: Extract text from PDF
        print("📖 Step 1: Reading CHUONG_V.pdf...")
        chuong_v_content = self.extract_text_from_pdf(chuong_v_file)
        if not chuong_v_content:
            return False
        
        # Save input text like your approach
        (self.process_folder / 'input.txt').write_text(chuong_v_content, encoding='utf-8')
        print(f"✅ Extracted {len(chuong_v_content)} characters")
        
        # Step 2: Process to markdown
        print("🔄 Step 2: Processing to markdown...")
        markdown_content = self.process_to_markdown(chuong_v_content)
        
        if markdown_content == "[KHÔNG TÌM THẤY]":
            print("❌ Failed to process to markdown")
            return False
        
        # Step 3: Convert markdown to DOCX
        print("📄 Step 3: Converting markdown to DOCX...")
        docx_path = self.markdown_to_docx(markdown_content)
        
        # Step 4: Replace placeholder using your proven method
        print("🔄 Step 4: Replacing placeholder in template...")
        if self.copy_template_to_output():
            doc = Document(self.output_file)
            
            # Use your proven placeholder replacement method
            self.replace_placeholder(doc, "{{muc_dich_cong_viec}}")
            
            doc.save(self.output_file)
            print(f"✅ SUCCESS: Full proven process completed!")
            print(f"📄 Check output file: {self.output_file}")
            return True
        
        return False

    @property
    def pdf_files(self):
        return {'CHUONG_V': 'CHUONG_V.pdf'}

def main():
    print("🇻🇳 Muc Dich Processor - Full Proven Process")
    print("=" * 50)
    
    try:
        processor = MucDichProcessor()
        print("✅ Environment loaded successfully")
        
        processor.test_muc_dich_cong_viec_full_process()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")

if __name__ == "__main__":
    main()