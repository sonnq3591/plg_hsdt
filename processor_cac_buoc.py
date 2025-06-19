import os
import shutil
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import fitz  # PyMuPDF
from dotenv import load_dotenv
import re
from copy import deepcopy

# Load environment variables
load_dotenv()

class CacBuocThucHienProcessor:
    def __init__(self):
        """Initialize the processor for {{cac_buoc_thuc_hien}} with precise PDF extraction"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("❌ OPENAI_API_KEY not found in .env file!")
        
        openai.api_key = self.openai_api_key
        
        self.pdf_folder = Path("pdf_inputs")
        self.template_file = "02_MUC_DO_HIEU_BIET_template.docx"
        self.output_file = "02_MUC_DO_HIEU_BIET_output.docx"
        
        # Create processing folder structure
        self.process_folder = Path("processed/cac_buoc_thuc_hien")
        self.process_folder.mkdir(parents=True, exist_ok=True)
        
        print("🎯 CacBuocThucHienProcessor initialized - EXACT PARAGRAPH TARGETING")

    def extract_text_from_pdf_precise(self, pdf_path):
        """Extract text with precise positioning using PyMuPDF"""
        try:
            doc = fitz.open(pdf_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                full_text += f"\n--- PAGE {page_num + 1} ---\n"
                full_text += text
            
            doc.close()
            return full_text
        except Exception as e:
            print(f"❌ Error reading PDF with PyMuPDF: {str(e)}")
            return None

    def extract_complete_section(self, pdf_content):
        """Extract EXACTLY what's visible in the PDF image"""
        prompt = f"""
Look at this PDF content and find section "5.2. Các bước thực hiện công việc:"

I can see from the PDF that there are exactly 2 paragraphs before the table:

1. First paragraph starts with: "- Không phân tán phông lưu trữ. Tài liệu của từng đơn vị hình thành phông phải được chỉnh lý và sắp xếp riêng biệt;"

2. Second paragraph starts with: "- Khi phân loại, lập hồ sơ (chỉnh sửa hoàn thiện, phục hồi hoặc lập mới hồ sơ), phải tôn trọng sự hình thành tài liệu theo trình tự theo dõi, giải quyết công việc;"

Extract these EXACT paragraphs as they appear, then extract the complete table.

Return format:

PARAGRAPH1: [Complete first paragraph starting with "- Không phân tán phông lưu trữ..."]
PARAGRAPH2: [Complete second paragraph starting with "- Khi phân loại, lập hồ sơ..."]
TABLE_START
Số TT,Nội dung công việc
1,Giao nhận tài liệu và lập biên bản giao nhận tài liệu
2,Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý (khoảng cách ~100m)
...continue with ALL rows including a), b), c)
TABLE_END

Extract EXACTLY what's written in the PDF.

PDF CONTENT:
{pdf_content}

EXTRACTED SECTION:
"""

        try:
            response = openai.ChatCompletion.create(
                model='gpt-4o',
                messages=[
                    {"role": "system", "content": "Extract EXACTLY what appears in the PDF. The user has identified the specific paragraphs that exist. Extract them word-for-word along with the complete table including all sub-rows a), b), c)."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4000,
                temperature=0.0
            )
            
            extracted_content = response.choices[0].message.content.strip()
            print("🎯 Extracted section content with dynamic detection")
            
            # Save for debugging
            (self.process_folder / 'extracted_complete_section.txt').write_text(extracted_content, encoding='utf-8')
            
            return extracted_content
            
        except Exception as e:
            print(f"❌ OpenAI API Error: {str(e)}")
            return None

    def parse_extracted_section(self, section_content):
        """Parse the extracted content - simplified for exact matching"""
        lines = section_content.split('\n')
        
        paragraph1 = ""
        paragraph2 = ""
        table_rows = []
        
        current_mode = "search"
        
        for line in lines:
            line = line.strip()
            
            if line.startswith('PARAGRAPH1:'):
                paragraph1 = line.replace('PARAGRAPH1:', '').strip()
            elif line.startswith('PARAGRAPH2:'):
                paragraph2 = line.replace('PARAGRAPH2:', '').strip()
            elif line == 'TABLE_START':
                current_mode = "table"
            elif line == 'TABLE_END':
                current_mode = "done"
            elif current_mode == "table" and line and ',' in line:
                table_rows.append(line)
        
        print(f"📄 Parsed: Para1({len(paragraph1)} chars), Para2({len(paragraph2)} chars), Table({len(table_rows)} rows)")
        
        # Show exactly what we extracted
        print(f"\n🔍 EXTRACTED CONTENT:")
        print(f"📄 Paragraph 1: '{paragraph1}'")
        print(f"📄 Paragraph 2: '{paragraph2}'")
        print(f"📋 Total table rows: {len(table_rows)}")
        
        # Show first 15 rows to see the sub-steps
        print("\n🔍 FIRST 15 TABLE ROWS:")
        for i, row in enumerate(table_rows[:15]):
            print(f"  Row {i+1}: {row}")
        
        return paragraph1, paragraph2, table_rows

    def is_sub_step(self, step_text):
        """Detect sub-steps for italic formatting"""
        if not step_text:
            return False
        
        step_clean = step_text.strip()
        
        # Pure numbers are main steps (NOT italic)
        if step_clean.isdigit():
            return False
        
        # Sub-step patterns that should be italic
        patterns = [
            r'^[a-z]\)$',      # a), b), c)
            r'^[a-z]$',        # a, b, c
            r'^[a-z]{2,}\)$',  # aa), bb), cc)
        ]
        
        for pattern in patterns:
            if re.match(pattern, step_clean):
                return True
                
        return True  # Everything else that's not a number

    def analyze_table_content(self, table_rows):
        """Analyze COMPLETE table content"""
        print(f"\n🔍 COMPLETE TABLE ANALYSIS ({len(table_rows)} total rows):")
        print("=" * 60)
        
        if not table_rows:
            print("❌ No table rows to analyze")
            return
        
        import csv
        from io import StringIO
        
        sub_step_count = 0
        main_step_count = 0
        
        for i, row_text in enumerate(table_rows):
            try:
                reader = csv.reader(StringIO(row_text))
                row = next(reader)
                
                if len(row) >= 2:
                    step_num = row[0].strip()
                    content = row[1].strip()
                    
                    is_substep = self.is_sub_step(step_num)
                    
                    if is_substep:
                        sub_step_count += 1
                        print(f"✨ Row {i+1}: '{step_num}' → ITALIC")
                    else:
                        main_step_count += 1
                        print(f"📝 Row {i+1}: '{step_num}' → Regular")
                    
            except Exception as e:
                print(f"❌ Error parsing row {i+1}: {e}")
        
        print(f"\n📊 SUMMARY: {main_step_count} main steps, {sub_step_count} sub-steps")
        if sub_step_count == 0:
            print("⚠️  WARNING: No sub-steps detected! Check extraction.")

    def create_formatted_docx(self, paragraph1, paragraph2, table_rows):
        """Create DOCX preserving complete table structure then applying formatting"""
        
        # First analyze what we have
        self.analyze_table_content(table_rows)
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # Add first paragraph
        if paragraph1:
            para = doc.add_paragraph(paragraph1)
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.15
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        # Add second paragraph
        if paragraph2:
            para = doc.add_paragraph(paragraph2)
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.15
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        # Create COMPLETE table preserving ALL rows
        if table_rows:
            import csv
            from io import StringIO
            
            parsed_rows = []
            for row_text in table_rows:
                try:
                    reader = csv.reader(StringIO(row_text))
                    row = next(reader)
                    parsed_rows.append(row)
                except:
                    continue

            if parsed_rows:
                # Create table with ALL rows
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                print(f"\n📋 Creating table with {len(parsed_rows)} total rows...")

                # Add ALL rows with formatting applied on top
                for i, row_data in enumerate(parsed_rows):
                    row = table.add_row().cells
                    
                    # Column 1: Số TT
                    cell1 = row[0]
                    cell1.width = Inches(1.0)
                    para1 = cell1.paragraphs[0]
                    para1.clear()
                    
                    step_number = str(row_data[0]).strip() if len(row_data) > 0 else ""
                    run1 = para1.add_run(step_number)
                    run1.font.name = 'Times New Roman'
                    run1.font.size = Pt(14)
                    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Column 2: Nội dung công việc
                    cell2 = row[1]
                    cell2.width = Inches(5.5)
                    para2 = cell2.paragraphs[0]
                    para2.clear()
                    content_text = str(row_data[1]) if len(row_data) > 1 else ""
                    run2 = para2.add_run(content_text)
                    
                    run2.font.name = 'Times New Roman'
                    run2.font.size = Pt(14)
                    
                    # Apply formatting on top of preserved structure
                    if self.is_sub_step(step_number):
                        run2.italic = True  # Content italic
                        run1.italic = True  # Step number italic
                        print(f"✨ Applied italic: '{step_number}' → {content_text[:30]}...")
                    
                    # Header row formatting
                    if i == 0:
                        run2.bold = True
                        run1.bold = True
                        run2.italic = False
                        run1.italic = False
                        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Save
        docx_path = self.process_folder / 'output.docx'
        doc.save(docx_path)
        print(f"✅ Created COMPLETE formatted DOCX: {docx_path}")
        
        return docx_path

    def replace_placeholder(self, doc, placeholder_tag):
        """Replace placeholder using proven method"""
        pattern = r"\{\{(.+?)\}\}"
        match = re.search(pattern, placeholder_tag)
        if not match:
            raise ValueError(f"Invalid placeholder: {placeholder_tag}")
        
        tag_name = match.group(1)
        source_path = Path(f"processed/{tag_name}/output.docx")

        if not source_path.exists():
            raise FileNotFoundError(f"Missing source doc: {source_path}")

        source_doc = Document(source_path)
        
        # Get all content
        all_elements = []
        for para in source_doc.paragraphs:
            if para.text.strip():
                all_elements.append(('paragraph', para))
        for table in source_doc.tables:
            all_elements.append(('table', table))

        # Replace placeholder
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)

                for element_type, element in reversed(all_elements):
                    if element_type == 'paragraph':
                        new_p = deepcopy(element._element)
                        parent.insert(index, new_p)
                    elif element_type == 'table':
                        new_t = deepcopy(element._element)
                        parent.insert(index, new_t)
                break

    def copy_template_to_output(self):
        """Copy template file to output file"""
        try:
            if not os.path.exists(self.template_file):
                print(f"❌ Template file not found: {self.template_file}")
                return False
            
            shutil.copy2(self.template_file, self.output_file)
            print(f"✅ Copied template to output file: {self.output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error copying template: {str(e)}")
            return False

    def test_complete_extraction(self):
        """Test COMPLETE extraction preserving all rows"""
        print("\n🧪 TESTING: EXACT PARAGRAPH TARGETING + CONSISTENT SUB-HEADINGS")
        print("=" * 60)
        
        # Check if CHUONG_V.pdf exists
        chuong_v_file = self.pdf_folder / 'CHUONG_V.pdf'
        if not chuong_v_file.exists():
            print(f"❌ File not found: {chuong_v_file}")
            return False
        
        # Step 1: Extract with PyMuPDF
        print("📖 Step 1: Reading PDF with PyMuPDF...")
        pdf_content = self.extract_text_from_pdf_precise(chuong_v_file)
        if not pdf_content:
            return False
        
        print(f"✅ Extracted {len(pdf_content)} characters")
        
        # Step 2: Exact paragraph targeting
        print("🔍 Step 2: Targeting exact paragraphs from PDF...")
        section_content = self.extract_complete_section(pdf_content)
        if not section_content:
            return False
        
        # Step 3: Parse exact content
        print("📋 Step 3: Parsing exact targeted content...")
        paragraph1, paragraph2, table_rows = self.parse_extracted_section(section_content)
        
        # Step 4: Create formatted DOCX
        print("📄 Step 4: Creating DOCX (preserve structure + apply formatting)...")
        try:
            docx_path = self.create_formatted_docx(paragraph1, paragraph2, table_rows)
        except Exception as e:
            print(f"❌ Failed to create DOCX: {str(e)}")
            return False
        
        # Step 5: Replace placeholder
        print("🔄 Step 5: Replacing placeholder...")
        if self.copy_template_to_output():
            doc = Document(self.output_file)
            self.replace_placeholder(doc, "{{cac_buoc_thuc_hien}}")
            doc.save(self.output_file)
            print("✅ SUCCESS: Exact paragraph targeting with consistent sub-headings!")
            print(f"📄 Check output file: {self.output_file}")
            return True
        
        return False

def main():
    print("🇻🇳 Cac Buoc Thuc Hien Processor - EXACT PARAGRAPH TARGETING")
    print("=" * 60)
    print("📋 STRATEGY: Target exact paragraphs + Consistent sub-heading formatting")
    print()
    
    try:
        processor = CacBuocThucHienProcessor()
        print("✅ Environment loaded successfully")
        
        processor.test_complete_extraction()
        
    except ValueError as e:
        print(e)
        print("💡 Please create a .env file with: OPENAI_API_KEY=your_key_here")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()